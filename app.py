"""
FPEL Proposal Generator — local web app
Run with: python -m streamlit run app.py
"""
import os, sys, tempfile, re, shutil
from datetime import datetime, date
import streamlit as st
import streamlit.components.v1 as components

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mappings import OPEX_MAPPINGS, CAPEX_MAPPINGS
from bom_transfer import transfer_values
from calculations import calc_environmental, calc_termination, fmt_indian, fmt_cap

# ── Backend-only helper: Rupees amount -> Indian-style words ──────────────────
# Used to auto-fill "In words: Rupees ___" from the Investment amount.
# Not exposed anywhere in the UI — purely computed at doc-generation time.
_W_ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
           "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
           "Seventeen", "Eighteen", "Nineteen"]
_W_TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]


def _w_two(n):
    if n < 20:
        return _W_ONES[n]
    return (_W_TENS[n // 10] + (" " + _W_ONES[n % 10] if n % 10 else "")).strip()


def _w_three(n):
    if n >= 100:
        rest = n % 100
        return _W_ONES[n // 100] + " Hundred" + (" " + _w_two(rest) if rest else "")
    return _w_two(n)


def amount_to_words_indian(rupees):
    """e.g. 25000000 -> 'Rupees Two Crore Fifty Lakh Only'"""
    rupees = int(round(rupees))
    if rupees <= 0:
        return ""
    crore, rupees = divmod(rupees, 10000000)
    lakh, rupees = divmod(rupees, 100000)
    thousand, hundred = divmod(rupees, 1000)
    parts = []
    if crore:
        parts.append(_w_three(crore) + " Crore")
    if lakh:
        parts.append(_w_three(lakh) + " Lakh")
    if thousand:
        parts.append(_w_three(thousand) + " Thousand")
    if hundred:
        parts.append(_w_three(hundred))
    return "Rupees " + " ".join(parts) + " Only"

st.set_page_config(
    page_title="FPEL · Proposal Generator",
    page_icon="https://raw.githubusercontent.com/twitter/twemoji/master/assets/svg/26a1.svg",
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items={
        "Get Help": None,
        "Report a bug": None,
        "About": "**FPEL Proposal Generator** · Fourth Partner Energy · Internal Tool"
    }
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Plus+Jakarta+Sans:wght@700;800;900&display=swap');

/* ── Force light mode ── */
html, body, #root, .stApp, [data-testid="stAppViewContainer"],
[data-testid="stHeader"], [data-testid="stToolbar"],
[data-baseweb="base-provider"], [data-baseweb="theme-provider"] {
  background-color: #F6F8FA !important;
  color-scheme: light !important;
}
[data-baseweb="input"], [data-baseweb="select"],
[data-baseweb="input"] > div, [data-baseweb="select"] > div {
  background: #FFFFFF !important;
  color: #1A1A2E !important;
}

:root {
  --orange:        #F7941D;
  --orange-dk:     #D4780A;
  --orange-lt:     #FFF4E6;
  --orange-md:     #FFD9A8;
  --orange-glow:   rgba(247,148,29,.20);
  --cyan:          #29ABE2;
  --cyan-dk:       #1A8BBF;
  --cyan-lt:       #E8F7FD;
  --cyan-md:       #A8DDF5;
  --cyan-glow:     rgba(41,171,226,.18);
  --grey:          #58595B;
  --grey-lt:       #F0F2F5;
  --grey-md:       #DDE1E7;
  --grey-dk:       #1A1A2E;
  --white:         #FFFFFF;
  --surface:       #F6F8FA;
  --card:          #FFFFFF;
  --text:          #1A1A2E;
  --text2:         #4A4A5A;
  --text3:         #8A8A9A;
  --border:        #E8EAF0;
  --border2:       #D0D4DC;
  --shadow-sm:     0 1px 4px rgba(0,0,0,.06);
  --shadow-md:     0 4px 16px rgba(0,0,0,.08);
  --shadow-lg:     0 8px 32px rgba(0,0,0,.10);
  --radius-sm:     8px;
  --radius-md:     12px;
  --radius-lg:     18px;
  --radius-xl:     24px;
}

/* ── Base ── */
.stApp, .main {
  background: var(--surface) !important;
  font-family: 'Inter', -apple-system, sans-serif !important;
}
.block-container {
  background: var(--surface) !important;
  padding-top: 1.5rem !important;
  max-width: 96% !important;
  width: 96% !important;
  padding-left: 2rem !important;
  padding-right: 2rem !important;
}
section[data-testid="stSidebar"] { display: none !important; }
#MainMenu, footer { visibility: hidden; }
[data-testid="stHeader"] { background: transparent !important; }

.stApp, .stApp p, .stApp span, .stApp li, .stApp div {
  font-family: 'Inter', -apple-system, sans-serif !important;
  color: var(--text) !important;
  line-height: 1.6 !important;
}

/* ── Top header bar ── */
.hdr {
  background: var(--white);
  border-radius: var(--radius-lg);
  padding: 1rem 1.5rem;
  margin-bottom: 1.25rem;
  display: flex;
  align-items: center;
  justify-content: space-between;
  box-shadow: var(--shadow-md);
  border: 1px solid var(--border);
  position: relative;
  overflow: hidden;
}
.hdr::before {
  content: '';
  position: absolute;
  left: 0; top: 0; bottom: 0;
  width: 5px;
  background: linear-gradient(180deg, var(--orange), var(--cyan));
  border-radius: 5px 0 0 5px;
}
.hdr-left { display: flex; align-items: center; gap: 1rem; }
.hdr-logo {
  width: 40px; height: 40px;
  background: var(--white);
  border-radius: 10px;
  display: flex; align-items: center; justify-content: center;
  border: 1.5px solid var(--border);
  overflow: hidden;
}
.hdr-logo img {
  width: 100%; height: 100%;
  object-fit: contain;
  padding: 4px;
}
.hdr h1 {
  font-family: 'Plus Jakarta Sans', sans-serif !important;
  font-size: 1rem;
  font-weight: 800;
  margin: 0;
  color: var(--grey-dk) !important;
  letter-spacing: -.02em;
}
.hdr h1 span.o { color: var(--orange) !important; }
.hdr h1 span.c { color: var(--cyan) !important; }
.hdr .sub {
  font-size: .68rem;
  color: var(--text3);
  margin-top: 1px;
  font-weight: 500;
  letter-spacing: .01em;
}
.badge {
  background: var(--grey-lt);
  color: var(--text3);
  border: 1px solid var(--border2);
  padding: .2rem .7rem;
  border-radius: 20px;
  font-size: .6rem;
  font-weight: 700;
  letter-spacing: .06em;
  font-family: 'Inter', sans-serif;
}

/* ── Step indicator ── */
.steps {
  display: flex; align-items: center; justify-content: center;
  gap: 0; margin: .75rem 0 1.25rem;
  background: var(--white);
  border-radius: var(--radius-md);
  padding: .85rem 1.5rem;
  box-shadow: var(--shadow-sm);
  border: 1px solid var(--border);
}
.step { display: flex; flex-direction: column; align-items: center; gap: 5px; }
.step-dot {
  width: 30px; height: 30px; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-family: 'Plus Jakarta Sans', sans-serif;
  font-size: .75rem; font-weight: 800;
  border: 2px solid var(--border2);
  color: var(--text3); background: var(--surface);
  transition: all .25s;
}
.step-dot.active {
  background: linear-gradient(135deg, var(--orange), #F5A94A);
  border-color: transparent;
  color: #fff;
  box-shadow: 0 4px 14px var(--orange-glow);
}
.step-dot.done {
  background: linear-gradient(135deg, var(--cyan), #3DC0F5);
  border-color: transparent;
  color: #fff;
}
.step-lbl { font-size: .6rem; font-weight: 700; color: var(--text3); text-transform: uppercase; letter-spacing: .05em; }
.step-lbl.active { color: var(--orange); font-weight: 800; }
.step-lbl.done  { color: var(--cyan); }
.step-line { width: 56px; height: 2px; background: var(--border2); margin: 0 6px; margin-bottom: 20px; }
.step-line.done { background: linear-gradient(90deg, var(--cyan), var(--cyan-md)); }

/* ── Hero / Landing ── */
.hero-wrap {
  background: linear-gradient(140deg, var(--white) 0%, #EEF8FD 50%, #FFF5EC 100%);
  border-radius: var(--radius-xl);
  padding: 4rem 2.5rem 3.5rem;
  text-align: center;
  margin: 0 0 1.5rem;
  border: 1px solid var(--border);
  box-shadow: var(--shadow-lg);
  position: relative;
  overflow: hidden;
}
.hero-wrap::before {
  content: '';
  position: absolute;
  top: -80px; right: -80px;
  width: 320px; height: 320px;
  background: radial-gradient(circle, rgba(247,148,29,.12) 0%, transparent 70%);
  border-radius: 50%;
}
.hero-wrap::after {
  content: '';
  position: absolute;
  bottom: -60px; left: -60px;
  width: 280px; height: 280px;
  background: radial-gradient(circle, rgba(41,171,226,.10) 0%, transparent 70%);
  border-radius: 50%;
}
.hero-wrap .over {
  font-size: .65rem; font-weight: 700; letter-spacing: .22em;
  text-transform: uppercase; color: var(--cyan) !important; margin-bottom: 1rem;
}
.hero-wrap h1 {
  font-family: 'Plus Jakarta Sans', sans-serif !important;
  font-size: 2.8rem; font-weight: 900;
  color: var(--grey-dk) !important;
  line-height: 1.08; margin-bottom: .85rem; letter-spacing: -.04em;
}
.hero-wrap h1 span { color: var(--orange) !important; }
.hero-wrap p {
  font-size: .88rem; color: var(--text2) !important;
  max-width: 420px; margin: 0 auto 2.5rem;
  line-height: 1.8; font-weight: 500;
}
.lcards {
  display: flex; gap: .85rem; justify-content: center;
  flex-wrap: wrap; margin: 1.5rem 0; position: relative; z-index: 1;
}
.lcard {
  background: var(--white);
  border: 1.5px solid var(--border);
  border-radius: var(--radius-md);
  padding: 1.2rem .9rem; width: 125px; text-align: center;
  transition: all .25s; box-shadow: var(--shadow-sm);
}
.lcard:hover {
  border-color: var(--orange);
  box-shadow: 0 6px 22px var(--orange-glow);
  transform: translateY(-5px);
}
.lcard .n {
  font-family: 'Plus Jakarta Sans', sans-serif !important;
  font-size: 1.55rem; font-weight: 900; color: var(--orange) !important; line-height: 1;
}
.lcard .l {
  font-size: .64rem; color: var(--text3) !important;
  margin-top: 6px; font-weight: 700; letter-spacing: .04em; text-transform: uppercase;
}
.foot { text-align: center; margin-top: 2rem; font-size: .66rem; color: var(--text3); font-weight: 500; }

/* ── Section label ── */
.sec {
  font-family: 'Inter', sans-serif;
  font-size: .62rem; font-weight: 700;
  color: var(--text3) !important;
  text-transform: uppercase; letter-spacing: .16em;
  padding: .25rem .75rem .25rem 0;
  margin: 1.75rem 0 .9rem;
  display: inline-flex; align-items: center; gap: .4rem;
}
.sec::before {
  content: '';
  display: inline-block; width: 14px; height: 2px;
  background: linear-gradient(90deg, var(--orange), var(--cyan));
  border-radius: 2px;
}
.oline { border: none; border-top: 1px solid var(--border); margin: 1.25rem 0; }
.ibox {
  background: var(--cyan-lt);
  border-left: 3px solid var(--cyan);
  padding: .65rem 1rem; border-radius: 0 var(--radius-sm) var(--radius-sm) 0;
  font-size: .79rem; color: var(--grey-dk) !important;
  margin: .75rem 0; font-weight: 600;
}
.pi-sec {
  font-size: .6rem; font-weight: 700; color: var(--orange) !important;
  text-transform: uppercase; letter-spacing: .14em;
}

/* ── Capacity hero card ── */
.cap {
  background: linear-gradient(135deg, var(--orange) 0%, #F5A94A 100%);
  padding: 1.5rem 2rem; border-radius: var(--radius-lg);
  text-align: center; margin: .8rem 0;
  box-shadow: 0 8px 30px var(--orange-glow);
  position: relative; overflow: hidden;
}
.cap::before {
  content: '';
  position: absolute; top: -30px; right: -30px;
  width: 120px; height: 120px;
  background: rgba(255,255,255,.10); border-radius: 50%;
}
.cap .v {
  font-family: 'Plus Jakarta Sans', sans-serif !important;
  font-size: 2.6rem; font-weight: 900; color: #fff !important; letter-spacing: -.04em;
}
.cap .l { font-size: .7rem; color: rgba(255,255,255,.7); margin-top: 4px; font-weight: 600; letter-spacing: .04em; }

/* ── Data cards ── */
.cr {
  background: var(--card);
  border: 1px solid var(--border);
  border-radius: var(--radius-md);
  padding: .75rem 1rem; margin: .25rem 0;
  display: flex; justify-content: space-between; align-items: center;
  transition: all .2s; box-shadow: var(--shadow-sm);
}
.cr:hover { border-color: var(--cyan-md); box-shadow: 0 4px 14px var(--cyan-glow); transform: translateY(-1px); }
.cr .lb { font-size: .68rem; color: var(--text3); font-weight: 600; text-transform: uppercase; letter-spacing: .05em; }
.cr .vl { font-family: 'Plus Jakarta Sans', sans-serif !important; font-size: .92rem; font-weight: 800; color: var(--grey-dk); }

.cr-hi {
  background: linear-gradient(135deg, var(--orange-lt), #FFF9F2);
  border: 1px solid var(--orange-md);
  border-radius: var(--radius-md);
  padding: .75rem 1rem; margin: .25rem 0;
  display: flex; justify-content: space-between; align-items: center;
  box-shadow: 0 2px 10px var(--orange-glow);
}
.cr-hi .lb { font-size: .68rem; color: var(--orange-dk); font-weight: 700; text-transform: uppercase; letter-spacing: .05em; }
.cr-hi .vl { font-family: 'Plus Jakarta Sans', sans-serif !important; font-size: .92rem; font-weight: 800; color: var(--grey-dk); }

/* ── Env cards ── */
.ec { border-radius: var(--radius-md); padding: 1.1rem 1rem; text-align: center; transition: all .22s; position: relative; overflow: hidden; }
.ec::before { content: ''; position: absolute; top: -18px; right: -18px; width: 64px; height: 64px; border-radius: 50%; opacity: .12; background: currentColor; }
.ec:hover { transform: translateY(-4px); }
.ec .v { font-family: 'Plus Jakarta Sans', sans-serif !important; font-size: 1.2rem; font-weight: 900; line-height: 1.1; }
.ec .l { font-size: .62rem; font-weight: 700; margin-top: 5px; text-transform: uppercase; letter-spacing: .05em; opacity: .75; }
.ec-coal { background: linear-gradient(135deg,#2C2C2E,#444446); border: none; box-shadow: 0 4px 16px rgba(44,44,46,.22); }
.ec-coal .v { color: #fff !important; } .ec-coal .l { color: rgba(255,255,255,.7) !important; }
.ec-co2  { background: linear-gradient(135deg,#F5A623,#F7C35A); border: none; box-shadow: 0 4px 16px rgba(245,166,35,.25); }
.ec-co2 .v  { color: #fff !important; } .ec-co2 .l  { color: rgba(255,255,255,.75) !important; }
.ec-water{ background: linear-gradient(135deg,#1A7FD4,#29ABE2); border: none; box-shadow: 0 4px 16px rgba(41,171,226,.25); }
.ec-water .v { color: #fff !important; } .ec-water .l { color: rgba(255,255,255,.75) !important; }
.ec-trees{ background: linear-gradient(135deg,#27AE60,#2ECC71); border: none; box-shadow: 0 4px 16px rgba(39,174,96,.22); }
.ec-trees .v { color: #fff !important; } .ec-trees .l { color: rgba(255,255,255,.75) !important; }

/* ── Upload page ── */
.upload-card {
  background: var(--card);
  border: 1.5px dashed var(--border2);
  border-radius: var(--radius-lg);
  padding: 1.4rem;
  transition: all .2s; box-shadow: var(--shadow-sm);
}
.upload-card:hover { border-color: var(--orange); box-shadow: 0 6px 22px var(--orange-glow); }
.upload-card-title {
  font-family: 'Inter', sans-serif; font-size: .7rem; font-weight: 700;
  color: var(--orange); text-transform: uppercase; letter-spacing: .1em;
  margin-bottom: .6rem; display: flex; align-items: center; gap: .4rem;
}
.pt-toggle { display: flex; gap: 0; background: var(--grey-lt); border-radius: var(--radius-md); padding: 4px; width: fit-content; margin-bottom: 1.25rem; }
.pt-opt { padding: .5rem 1.6rem; border-radius: var(--radius-sm); font-family: 'Inter', sans-serif; font-size: .85rem; font-weight: 700; cursor: pointer; transition: all .2s; color: var(--text3); border: none; background: transparent; }
.pt-opt.active { background: var(--white); color: var(--orange); box-shadow: var(--shadow-md); }

/* ── File uploaders ── */
.stFileUploader > div { border: 1.5px dashed var(--border2) !important; border-radius: var(--radius-md) !important; background: var(--grey-lt) !important; transition: all .2s !important; }
.stFileUploader > div:hover { border-color: var(--orange) !important; background: var(--orange-lt) !important; }
.stFileUploader p, .stFileUploader span { color: var(--text2) !important; font-weight: 600 !important; }

/* ── Form labels ── */
.stTextInput > label,
.stNumberInput > label,
.stSelectbox > label,
.stFileUploader > label,
.stTextArea > label {
  font-family: 'Inter', sans-serif !important;
  font-size: .69rem !important; font-weight: 600 !important;
  color: var(--text2) !important;
  text-transform: uppercase !important; letter-spacing: .05em !important;
  margin-bottom: .25rem !important;
}

/* ── Text / Number inputs ── */
.stTextInput > div > div,
.stNumberInput > div > div,
.stTextArea > div > div {
  border: 1.5px solid var(--border2) !important;
  border-radius: var(--radius-sm) !important;
  background: var(--white) !important;
  box-shadow: var(--shadow-sm) !important;
  transition: all .2s !important;
}
.stTextInput > div > div:focus-within,
.stNumberInput > div > div:focus-within,
.stTextArea > div > div:focus-within {
  border-color: var(--orange) !important;
  box-shadow: 0 0 0 3px var(--orange-glow), var(--shadow-sm) !important;
}
.stTextInput input, .stNumberInput input, .stTextArea textarea {
  color: var(--text) !important; background: transparent !important;
  font-family: 'Inter', sans-serif !important;
  font-size: .88rem !important; font-weight: 600 !important;
  caret-color: var(--orange) !important;
}
.stTextInput input::placeholder, .stNumberInput input::placeholder { color: var(--border2) !important; font-weight: 400 !important; }
.stNumberInput button { background: var(--grey-lt) !important; border: none !important; color: var(--orange) !important; font-weight: 800 !important; }
.stNumberInput button:hover { background: var(--orange-lt) !important; }

/* ── Selectbox ── */
.stSelectbox > div > div { border: 1.5px solid var(--border2) !important; border-radius: var(--radius-sm) !important; background: var(--white) !important; box-shadow: var(--shadow-sm) !important; transition: all .2s !important; }
.stSelectbox > div > div:focus-within { border-color: var(--orange) !important; box-shadow: 0 0 0 3px var(--orange-glow) !important; }
.stSelectbox [data-baseweb="select"] *, .stSelectbox [data-baseweb="select"] span, .stSelectbox [data-baseweb="select"] div, .stSelectbox [data-baseweb="select"] p { color: var(--text) !important; background: transparent !important; font-weight: 600 !important; font-family: 'Inter', sans-serif !important; }
.stSelectbox svg { color: var(--orange) !important; fill: var(--orange) !important; opacity: 1 !important; }

/* ── Dropdown popup ── */
[data-baseweb="popover"], [data-baseweb="popover"] > div,
[data-baseweb="popover"] [data-baseweb="menu"], [data-baseweb="menu"], ul[data-baseweb="menu"] {
  background: var(--white) !important; border: 1px solid var(--border2) !important;
  border-radius: var(--radius-md) !important; box-shadow: var(--shadow-lg) !important; padding: 4px !important;
}
[data-baseweb="popover"] li, [data-baseweb="menu"] li, [role="option"], [data-baseweb="menu-item"] {
  background: var(--white) !important; color: var(--text) !important;
  font-size: .84rem !important; font-weight: 500 !important;
  font-family: 'Inter', sans-serif !important;
  border-radius: var(--radius-sm) !important; padding: .45rem .8rem !important;
}
[data-baseweb="popover"] li:hover, [data-baseweb="menu"] li:hover, [role="option"]:hover, [aria-selected="true"] { background: var(--orange-lt) !important; color: var(--orange) !important; }
[data-baseweb="popover"] * { color: var(--text) !important; }
[data-baseweb="popover"] [aria-selected="true"], [data-baseweb="popover"] [aria-selected="true"] * { color: var(--orange) !important; }

/* ── Buttons ── */
.stButton > button[kind="primary"] {
  background: linear-gradient(135deg, var(--orange) 0%, #F5A94A 100%) !important;
  color: #fff !important; border: none !important; font-weight: 700 !important;
  border-radius: var(--radius-sm) !important; padding: .6rem 1.5rem !important;
  box-shadow: 0 4px 16px var(--orange-glow) !important;
  transition: all .2s !important; font-family: 'Inter', sans-serif !important;
  font-size: .86rem !important; letter-spacing: .01em !important;
}
.stButton > button[kind="primary"]:hover { transform: translateY(-2px) !important; box-shadow: 0 8px 24px var(--orange-glow) !important; }
.stButton > button[kind="secondary"] {
  border: 1.5px solid var(--border2) !important; color: var(--text2) !important;
  background: var(--white) !important; border-radius: var(--radius-sm) !important;
  font-weight: 600 !important; transition: all .2s !important;
  font-family: 'Inter', sans-serif !important; font-size: .86rem !important;
  box-shadow: var(--shadow-sm) !important;
}
.stButton > button[kind="secondary"]:hover { border-color: var(--orange) !important; color: var(--orange) !important; background: var(--orange-lt) !important; }
.stDownloadButton > button {
  background: linear-gradient(135deg, var(--cyan), #3DC0F5) !important;
  color: #fff !important; border: none !important; border-radius: var(--radius-sm) !important;
  font-weight: 700 !important; font-family: 'Inter', sans-serif !important;
  box-shadow: 0 4px 16px var(--cyan-glow) !important;
}

/* ── Progress ── */
.stProgress > div > div > div { background: linear-gradient(90deg, var(--orange), var(--cyan)) !important; border-radius: 8px !important; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] { gap: 0; border-bottom: 2px solid var(--border); background: transparent; }
.stTabs [data-baseweb="tab"] {
  color: var(--text3) !important; font-weight: 600; border-bottom: 3px solid transparent;
  padding: .5rem 1rem; font-size: .78rem; font-family: 'Inter', sans-serif !important;
  letter-spacing: .01em; background: transparent !important; transition: color .15s !important;
}
.stTabs [aria-selected="true"] { color: var(--orange) !important; font-weight: 700; border-bottom: 3px solid var(--orange) !important; }
.stTabs [data-baseweb="tab"]:hover { color: var(--orange) !important; }

/* ── Expanders ── */
.streamlit-expanderHeader { color: var(--text) !important; font-weight: 600 !important; font-size: .82rem !important; font-family: 'Inter', sans-serif !important; background: var(--grey-lt) !important; border-radius: var(--radius-sm) !important; border: 1px solid var(--border) !important; }
.streamlit-expanderContent { background: var(--white) !important; border: 1px solid var(--border) !important; border-top: none !important; border-radius: 0 0 var(--radius-sm) var(--radius-sm) !important; }

/* ── Alerts ── */
div[data-testid="stAlert"] { border-radius: var(--radius-md) !important; }

/* ── Headings / Captions ── */
.stApp h3 { font-family: 'Plus Jakarta Sans', sans-serif !important; font-size: 1rem !important; font-weight: 800 !important; color: var(--grey-dk) !important; }
.stApp .stCaption { font-size: .72rem !important; color: var(--text3) !important; font-weight: 500 !important; }

/* ── Help button (styling moved to the popover block near its usage,
      so both the native st.popover version and the fallback share one
      place) ── */

/* ── Radio (project type) ── */
div[data-testid="stRadio"] > label { display: none !important; }
div[data-testid="stRadio"] > div { background: transparent !important; gap: .5rem !important; }
div[data-testid="stRadio"] > div > label { background: var(--grey-lt) !important; border: 1.5px solid var(--border2) !important; border-radius: var(--radius-sm) !important; padding: .4rem 1rem !important; font-weight: 600 !important; color: var(--text2) !important; cursor: pointer !important; transition: all .2s !important; font-size: .83rem !important; }
div[data-testid="stRadio"] > div > label:has(input:checked) { background: var(--orange-lt) !important; border-color: var(--orange) !important; color: var(--orange) !important; }
div[data-testid="stRadio"] input { display: none !important; }

/* ── Upload col ── */
[data-testid="column"] .upload-card-title { margin-top: .25rem; }
</style>""", unsafe_allow_html=True)

# ── Persistent UI: Save + Scroll ──
# The save button serialises the current session state to a JSON draft file
# which the user can download and re-upload later to restore exactly where they left off.
# We use a hidden Streamlit form so clicking Save triggers a rerun that produces
# the download — pure JS localStorage is blocked inside the Claude.ai iframe.

_DRAFT_KEYS = [
    "project_type", "prefilled", "key_inputs", "stage",
    "site_image_path", "final_input_path", "tmp_dir",
    "output_path", "output_name",
]

# ── Save: write session → JSON download ──
if st.session_state.get("_save_triggered"):
    import json as _json
    _draft = {}
    for _k in _DRAFT_KEYS:
        _v = st.session_state.get(_k)
        try:
            _json.dumps(_v)          # only include JSON-serialisable values
            _draft[_k] = _v
        except Exception:
            pass
    # also capture every f_PH_ widget value currently live
    _draft["_widgets"] = {
        k: v for k, v in st.session_state.items()
        if k.startswith("f_") or k.startswith("ki_") or k.startswith("fa9_")
    }
    _draft_bytes = _json.dumps(_draft, ensure_ascii=False, indent=2).encode("utf-8")
    st.session_state["_save_triggered"] = False
    st.session_state["_draft_bytes"] = _draft_bytes

# ── Load: restore JSON → session state ──
if st.session_state.get("_load_triggered") and st.session_state.get("_draft_upload"):
    import json as _json
    try:
        _loaded = _json.loads(st.session_state["_draft_upload"].getvalue())
        for _k, _v in _loaded.items():
            if _k == "_widgets":
                for _wk, _wv in _v.items():
                    st.session_state[_wk] = _wv
            else:
                st.session_state[_k] = _v
        st.session_state["_load_triggered"] = False
        st.session_state["_draft_upload"] = None
        st.success("Draft restored. Continue from where you left off.")
        st.rerun()
    except Exception as _le:
        st.error(f"Could not restore draft: {_le}")
    st.session_state["_load_triggered"] = False

# Save/load removed per user request

STAGE_MAP = {"landing":0,"upload":1,"key_inputs":2,"form":3,"generate":4,"done":4}

def render_live_timer(stage):
    """Small stopwatch badge — ticks live from the 'Get Started' click, freezes
    on the Done page showing the final time taken to generate the proposal."""
    started = st.session_state.get("_proposal_started_at")
    if not started:
        return
    if stage == "done":
        frozen = st.session_state.get("_last_time_taken", "—")
        components.html(f"""
        <div style="display:flex;justify-content:flex-end;font-family:'Courier New',monospace;">
          <div style="background:#000;color:#fff;font-weight:800;
            padding:.4rem .9rem;border-radius:20px;font-size:.85rem;display:inline-block;">
            ⏱ {frozen} (finished)
          </div>
        </div>""", height=42)
    else:
        _start_ms = int(started * 1000)
        components.html(f"""
        <div style="display:flex;justify-content:flex-end;font-family:'Courier New',monospace;">
          <div id="fpel-timer" style="background:#000;color:#fff;font-weight:800;
            padding:.4rem .9rem;border-radius:20px;font-size:.85rem;display:inline-block;">
            ⏱ 00:00
          </div>
        </div>
        <script>
        (function(){{
          const startMs = {_start_ms};
          const el = document.getElementById('fpel-timer');
          function tick(){{
            const diff = Math.max(0, Math.floor((Date.now() - startMs)/1000));
            const h = Math.floor(diff/3600);
            const m = Math.floor((diff%3600)/60).toString().padStart(2,'0');
            const s = (diff%60).toString().padStart(2,'0');
            if (el) el.textContent = '⏱ ' + (h>0 ? h+':' : '') + m + ':' + s;
          }}
          tick();
          setInterval(tick, 1000);
        }})();
        </script>
        """, height=42)

@st.cache_data
def _logo_b64():
    with open(os.path.join(os.path.dirname(__file__), "logo.png"), "rb") as f:
        import base64
        return base64.b64encode(f.read()).decode()

def show_header(stage=""):
    idx = STAGE_MAP.get(stage, 0)
    steps = [("Upload","upload"),("Key Inputs","key_inputs"),("Review","form"),("Generate","done")]
    dots_html = ""
    for i,(lbl,s) in enumerate(steps):
        cls = "done" if idx > STAGE_MAP.get(s,99) else ("active" if idx == STAGE_MAP.get(s,99) else "")
        dots_html += f'<div class="step"><div class="step-dot {cls}">{i+1}</div><div class="step-lbl {cls}">{lbl}</div></div>'
        if i < len(steps)-1:
            line_cls = "done" if idx > STAGE_MAP.get(s,99) else ""
            dots_html += f'<div class="step-line {line_cls}"></div>'

    st.markdown(f"""
    <div class="hdr">
      <div class="hdr-left">
        <div class="hdr-logo"><img src="data:image/png;base64,{_logo_b64()}" alt="FPEL logo"></div>
        <div>
          <h1><span class="c">FOURTH</span>&nbsp;<span class="c">PARTNER</span>&nbsp;<span class="o">ENERGY</span></h1>
          <div class="sub">Proposal Generator &nbsp;·&nbsp; Solar & Clean Energy Solutions</div>
        </div>
      </div>
      <div class="badge">INTERNAL TOOL</div>
    </div>
    <div class="steps">{dots_html}</div>
    """, unsafe_allow_html=True)
    render_live_timer(stage)

APP_DIR = os.path.dirname(os.path.abspath(__file__))
PROPOSAL_LOG_FILE = os.path.join(APP_DIR, "proposal_log.csv")

PROPOSAL_LOG_HEADER = ["Date","Submitted By","Capacity (kWp)","Project Type","File Name","Time Taken"]

def _migrate_proposal_log():
    """One-time fix for log files written before the 'Time Taken' column
    existed — their header row is missing it, which makes csv.DictReader
    misalign every later 6-value row (the 6th value lands under the special
    None 'restkey' instead of 'Time Taken', so it never showed up)."""
    import csv
    if not os.path.exists(PROPOSAL_LOG_FILE):
        return
    try:
        with open(PROPOSAL_LOG_FILE, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)
        if not rows or rows[0] == PROPOSAL_LOG_HEADER:
            return   # already current
        old_header = rows[0]
        n = len(PROPOSAL_LOG_HEADER)
        fixed = [PROPOSAL_LOG_HEADER]
        for r in rows[1:]:
            r = r + ["—"] * (n - len(r))   # pad short rows (old schema) with —
            fixed.append(r[:n])
        with open(PROPOSAL_LOG_FILE, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerows(fixed)
    except Exception:
        pass

def log_proposal(submitter, capacity_kwp, project_type, filename, time_taken=""):
    """Append one row to the local proposal generation log (CSV)."""
    import csv
    _migrate_proposal_log()
    is_new = not os.path.exists(PROPOSAL_LOG_FILE)
    try:
        with open(PROPOSAL_LOG_FILE, "a", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            if is_new:
                w.writerow(PROPOSAL_LOG_HEADER)
            w.writerow([datetime.now().strftime("%d-%m-%Y %H:%M"), submitter or "—",
                        capacity_kwp or "—", project_type or "—", filename or "—", time_taken or "—"])
    except Exception:
        pass   # logging must never block proposal generation

def read_proposal_log(limit=25):
    """Return the most recent `limit` log rows, newest first. Pass limit=None for all rows."""
    import csv
    if not os.path.exists(PROPOSAL_LOG_FILE):
        return []
    _migrate_proposal_log()
    try:
        with open(PROPOSAL_LOG_FILE, newline="", encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
        rows = list(reversed(rows))
        return rows if limit is None else rows[:limit]
    except Exception:
        return []

def format_elapsed(seconds):
    """Format a duration in seconds as e.g. '2m 14s' or '45s'."""
    try:
        seconds = int(round(float(seconds)))
    except Exception:
        return "—"
    m, s = divmod(max(seconds, 0), 60)
    h, m = divmod(m, 60)
    if h: return f"{h}h {m}m {s}s"
    if m: return f"{m}m {s}s"
    return f"{s}s"
LOCAL_OPEX_TEMPLATE = None
for n in ["OPEX_Proposal_-_Final.docx","OPEX_Proposal_-_Final__1_.docx","OPEX Proposal - Final.docx"]:
    p = os.path.join(APP_DIR, n)
    if os.path.exists(p): LOCAL_OPEX_TEMPLATE = p; break

LOCAL_CAPEX_TEMPLATE = None
for n in ["CAPEX_Proposal_-_Coding.docx","CAPEX_Proposal_-_Final.docx","CAPEX Proposal - Final.docx"]:
    p = os.path.join(APP_DIR, n)
    if os.path.exists(p): LOCAL_CAPEX_TEMPLATE = p; break

defaults = {"stage":"landing","transfer_results":[],"transfer_warnings":[],"tmp_dir":None,
    "site_image_path":None,"updated_proposal_path":None,"word_template_path":None,
    "final_input_path":None,"output_path":None,"output_name":None,"prefilled":{},"n_missing":0,
    "capacity_kwp":0.0,"key_inputs":{},"project_type":"OPEX"}
for k,v in defaults.items():
    if k not in st.session_state: st.session_state[k]=v

SECTIONS = {

    # ════════ SHARED SECTIONS (PH_001–084) — same in both CAPEX and OPEX ════════

    "1 · Cover & Executive Summary (CAPEX)": [
        ("PH_001","Proposed Capacity (kWp)","A"),
        ("PH_002","Installation Technology","DD_INSTTECH"),
        ("PH_003","Estimated Generation (kWh/year)","A"),
        ("PH_004","% Consumption Replaced","M"),
        ("PH_005","Net Metering","DD_NETMETER"),
        ("PH_006","Key Components — Modules","M"),
        ("PH_007","Key Components — Inverters","M"),
        ("PH_008","Project Timeline","DD_TIMELINE"),
        ("PH_009","Exclusions","TEXTAREA"),
        ("PH_010","Investment (₹ Cr)","A"),
        ("PH_011","Payment Terms","TEXTAREA"),
        ("PH_012","AMC Cost","A"),
        ("PH_013","Payback Period","A"),
        ("PH_014","Net Savings over 25 Years","A"),
        ("PH_015","CO₂ Reduction (Tonnes/year)","A"),
        ("PH_016","Submitted by","DD_SUBMITTER"),
        ("PH_017","Submitted on","DATE"),
        ("PH_018","Validity","O"),
    ],
    "1 · Cover & Executive Summary (OPEX)": [
        ("PH_001","Proposed Capacity (kWp)","A"),
        ("PH_002","Installation Technology","DD_INSTTECH"),
        ("PH_003","Estimated Generation (kWh/year)","A"),
        ("PH_004","Guaranteed Generation (kWh/year)","A"),
        ("PH_005","% Consumption Replaced","A"),
        ("PH_006","Net Metering (Yes/No)","YN"),
        ("PH_007","Key Components — Modules","M"),
        ("PH_008","Key Components — Inverters","M"),
        ("PH_009","Project Timeline","DD_TIMELINE"),
        ("PH_010","Exclusions","TEXTAREA"),
        ("PH_011","Variable Grid Rate (₹/kWh)","A"),
        ("PH_012","Proposed Solar Tariff (₹/kWh)","A"),
        ("PH_013","% Savings (1st Year)","A"),
        ("PH_014","% Grid Consumption Replaced","A"),
        ("PH_015","CO₂ Reduction (Tons/year)","A"),
        ("PH_016","Submitted by","DD_SUBMITTER"),
        ("PH_017","Submitted on","DATE"),
        ("PH_166","Annual Electricity Consumption (kWh/year)","M"),
        ("PH_167","% Consumption Replaced (auto)","A"),
    ],

    "2 · Site Details": [
        ("PH_019","Customer Name","M"),
        ("PH_020","Customer Address (short)","M"),
        ("PH_021","Contact Person","M"),
        ("PH_022","Mobile","M"),
        ("PH_023","Email","M"),
        ("PH_024","Site Address (full)","M"),
        ("PH_025","Latitude","M"),
        ("PH_026","Longitude","M"),
        ("PH_027","Type of Installation","DD_INST"),
        ("PH_028","Type of Roofs","DD_ROOF"),
        ("PH_029","Availability of Water","YN"),
        ("PH_030","Shading","YN_NO"),
    ],

    "3 · Capacity Assessment": [
        ("PH_032","Shadow-free area (kWp)","M"),
        ("PH_033","Shadow-free area — Notes","O"),
        ("PH_034","Electricity consumption (kWp)","M"),
        ("PH_035","Electricity consumption — Notes","O"),
        ("PH_036","State regulatory limit (kWp)","O"),
        ("PH_037","State regulatory limit — Notes","O"),
        ("PH_038","Proposed Capacity (kWp)","A"),
        ("PH_039","Proposed Capacity — Notes","O"),
    ],

    "4 · Plant Layout": [
        ("PH_040","Type of Installation","DD_INST"),
        ("PH_041","Type of Mounting","M"),
        ("PH_042","Power Evacuation","DD_EVAC"),
        ("PH_043","Rooftop Row1 — Shed/Building","O"),
        ("PH_044","Rooftop Row1 — Capacity (kWp)","O"),
        ("PH_045","Rooftop Row1 — Tilt","O"),
        ("PH_046","Rooftop Row1 — Azimuth","O"),
        ("PH_047","Rooftop Row1 — Walkways/Lifeline","O"),
        ("PH_048","Rooftop Row2 — Shed/Building","O"),
        ("PH_049","Rooftop Row2 — Capacity (kWp)","O"),
        ("PH_050","Rooftop Row2 — Tilt","O"),
        ("PH_051","Rooftop Row2 — Azimuth","O"),
        ("PH_052","Rooftop Row2 — Walkways/Lifeline","O"),
        ("PH_053","Rooftop Total Capacity (kWp)","A"),
        ("PH_054","Shed Row1 — Shed/Building","O"),
        ("PH_055","Shed Row1 — Capacity (kWp)","O"),
        ("PH_056","Shed Row1 — Tilt","O"),
        ("PH_057","Shed Row1 — Azimuth","O"),
        ("PH_058","Shed Row1 — Walkways/Railings","O"),
        ("PH_059","Shed Row2 — Shed/Building","O"),
        ("PH_060","Shed Row2 — Capacity (kWp)","O"),
        ("PH_061","Shed Row2 — Tilt","O"),
        ("PH_062","Shed Row2 — Azimuth","O"),
        ("PH_063","Shed Row2 — Walkways/Railings","O"),
        ("PH_064","Shed Total Capacity (kWp)","O"),
    ],

    "5 · Estimated Generation (CAPEX)": [
        ("PH_079","Type of System","DD_SYSTYPE"),
        ("PH_080","Estimated Generation (kWh/year)","A"),
        ("PH_081","Degradation","O"),
        ("PH_082","Guaranteed Generation (kWh/year)","A"),
        ("PH_083","Plant Life (Years)","M"),
        ("PH_084","GHI / Irradiation (kWh/m²/yr)","M"),
    ],

    "5 · Estimated Generation (OPEX)": [
        ("PH_065","Type of System","DD_SYSTYPE"),
        ("PH_066","Estimated Generation (kWh/year)","A"),
        ("PH_067","Degradation","O"),
        ("PH_068","Guaranteed Generation (kWh/year)","A"),
        ("PH_069","Plant Life (Years)","M"),
        ("PH_070","GHI / Irradiation (kWh/m²/yr)","M"),
    ],

    "6 · Bill of Material (CAPEX)": [
        ("PH_085","PV Module — Type","M"),
        ("PH_086","PV Module — Make/Model","M"),
        ("PH_087","PV Module — Efficiency","O"),
        ("PH_088","Inverter — Type","M"),
        ("PH_089","Inverter — Env.Class/Location","O"),
        ("PH_090","Inverter — Make/Model","M"),
        ("PH_091","Mounting Structure — Type","M"),
        ("PH_092","Mounting Structure — Make","O"),
        ("PH_093","DC Cable — Make/Size","O"),
        ("PH_094","AC Cable — Make/Size","O"),
        ("PH_095","Cable Tray — Type/Size","O"),
        ("PH_096","LT Switchgear — Type/Make/Rating","O"),
        ("PH_097","Spare Feeder — Type/Make/Rating","O"),
        ("PH_098","Energy Meter + ACCB — Make/Class","O"),
        ("PH_099","Lightning Arrestor — Type/Make","O"),
        ("PH_100","Earthing Kit — Make/Specs","O"),
        ("PH_101","Lifeline — Make","O"),
        ("PH_102","Walkways — Type/Size","O"),
        ("PH_103","Mesh over Skylights — Make/Specs","O"),
        ("PH_104","Railings over Roof — Make/Specs","O"),
        ("PH_105","Ladder — Make/Specs","O"),
        ("PH_106","Remote Monitoring System","O"),
        ("PH_107","Irradiation Sensor — Make/Specs","O"),
        ("PH_108","Ambient Temp. Sensor","O"),
        ("PH_109","Module Cleaning (Manual)","O"),
        ("PH_110","Water Pump/Tank/Meter","O"),
        ("PH_111","HT/VCB Panel","O"),
        ("PH_112","Transformer","O"),
        ("PH_113","Transmission/Termination","O"),
        ("PH_114","Auxiliary Power","O"),
        ("PH_115","MCR Room / SCADA","O"),
        ("PH_116","Boundary/Lighting/CCTV","O"),
        ("PH_117","Fire Fighting Equipment","O"),
        ("PH_118","Net Metering Hardware","O"),
        ("PH_119","Net Metering Cubicle","O"),
    ],

    "6 · Bill of Material (OPEX)": [
        ("PH_071","PV Module — Type","M"),
        ("PH_072","PV Module — Make/Model","M"),
        ("PH_073","PV Module — Efficiency","O"),
        ("PH_074","Inverter — Type","M"),
        ("PH_075","Inverter — Env.Class/Location","O"),
        ("PH_076","Inverter — Make/Model","M"),
        ("PH_077","Mounting Structure — Type","M"),
        ("PH_078","Mounting Structure — Make","O"),
        ("PH_079","DC Cable — Make/Size","O"),
        ("PH_080","AC Cable — Make/Size","O"),
        ("PH_081","Cable Tray — Type/Size","O"),
        ("PH_082","LT Switchgear — Type/Make/Rating","O"),
        ("PH_083","Spare Feeder — Type/Make/Rating","O"),
        ("PH_084","Energy Meter + ACCB — Make/Class","O"),
        ("PH_085","Lightning Arrestor — Type/Make","O"),
        ("PH_086","Earthing Kit — Make/Specs","O"),
        ("PH_087","Lifeline — Make","O"),
        ("PH_088","Walkways — Type/Size","O"),
        ("PH_089","Mesh over Skylights — Make/Specs","O"),
        ("PH_090","Railings over Roof — Make/Specs","O"),
        ("PH_091","Ladder — Make/Specs","O"),
        ("PH_092","Remote Monitoring System","O"),
        ("PH_093","Irradiation Sensor — Make/Specs","O"),
        ("PH_094","Ambient Temp. Sensor","O"),
        ("PH_095","Module Cleaning (Manual)","O"),
        ("PH_096","Water Pump/Tank/Meter","O"),
        ("PH_097","Net Metering Hardware","O"),
        ("PH_098","Zero Feed in Device","O"),
        ("PH_099","DG Synchronization","O"),
    ],

    # ════════ CAPEX-ONLY SECTIONS ════════

    "7 · Project Schedule (CAPEX)": [
        ("PH_120","Delivery & Installation","DD_TIMELINE"),
    ],
    "8 · Project Cost": [
        ("PH_121","EPC — Investment per Wp (₹)","M"),
        ("PH_122","EPC — Total Investment (₹)","A"),
        ("PH_123","Net Metering — Investment per Wp (₹)","O"),
        ("PH_124","Net Metering — Total Investment (₹)","O"),
        ("PH_125","Generation Panel — Investment per Wp (₹)","O"),
        ("PH_126","Generation Panel — Total Investment (₹)","O"),
        ("PH_127","TOTAL — Investment per Wp (₹)","A"),
        ("PH_128","TOTAL — Total Investment (₹) with GST of 8.9%","A"),
        ("PH_129","Total in words","O"),
        ("PH_130","USD–INR Exchange Rate","O"),
        ("PH_131","AMC Cost Year 1 (₹/kWp)","O"),
    ],
    "9 · Financial Analysis": [
        ("PH_132","System Size (kWp)","A"),
        ("PH_133","System Cost incl. GST (₹ Lacs)","A"),
        ("PH_134","GST Input Credit (₹ Lacs)","A"),
        ("PH_135","Net Cost to Client (₹ Lacs)","A"),
        ("PH_136","AMC Cost (₹ Lacs/year)","A"),
        ("PH_137","Solar Units Generated Yr1","A"),
        ("PH_138","Present Power Tariff (₹/unit)","A"),
        ("PH_139","Avg. EB Tariff Increase (% p.a.)","A"),
        ("PH_140","Savings in Year 1 post-tax (₹ Lacs)","A"),
        ("PH_141","Payback Period (Years)","A"),
        ("PH_142","Project Life (Years)","A"),
        ("PH_143","Inverter Life (Years)","A"),
        ("PH_144","Net Savings over Project Life (₹ Lacs)","A"),
        ("PH_145","Total Units over Project Life (Lac Units)","A"),
        ("PH_146","Project IRR Post Tax (%)","A"),
        ("PH_147","Equity IRR Post Tax (%)","A"),
        ("PH_148","Levelised Cost of Generation (₹/kWh)","A"),
    ],
    "10 · Terms & Conditions (CAPEX)": [
        ("PH_149","Offer Validity","O"),
        ("PH_150","AMC Cost Yr1 (₹ Lakhs)","A"),
        ("PH_151","FPEL GST No. (Location)","DD_GST"),
        ("PH_152","Exclusions (T&C)","TEXTAREA"),
    ],

    # ════════ OPEX-ONLY SECTIONS ════════

    "7 · Project Schedule (OPEX)": [
        ("PH_100","Delivery & Installation","DD_TIMELINE"),
    ],
    "8 · Commercials (OPEX)": [
        ("PH_101","Offtaker — Legal Name","M"),
        ("PH_102","Offtaker — Registered Office","M"),
        ("PH_103","Contract Term (Years)","DD_CONTRACT"),
        ("PH_104","Lock-in Period (Years)","M"),
        ("PH_105","Tariff — Period","M"),
        ("PH_106","Tariff — Rate (₹/kWh)","M"),
        ("PH_107","Tariff — Type","DD_TARIFF"),
        ("PH_108","USD–INR Exchange Rate","M"),
        ("PH_109","Variable Grid Rate (₹/kWh)","A"),
        ("PH_110","Proposed Solar Tariff (₹/kWh)","A"),
        ("PH_111","Landed Cost of Solar Power (₹/kWh)","O"),
        ("PH_112","Per Unit Savings (₹/kWh)","A"),
        ("PH_113","Annual Generation Yr1 (kWh)","A"),
        ("PH_114","1st Yr Savings (₹)","A"),
        ("PH_115","Total Savings over PPA Term (₹)","A"),
        ("PH_116","Total Savings in words","A"),
    ],
    "11 · Termination Charges (OPEX)": [
        (f"PH_{130+i}", f"Year {i+1}", "A") for i in range(25)
    ],
}

ALL_PH = [ph for fields in SECTIONS.values() for ph,_,_ in fields]

TIMELINE_OPTS  = ["","3 + 1 Month","4 + 1 Month","5 + 1 Month","6 + 1 Month"]
CONTRACT_OPTS  = ["","10","15","20","25"]
TARIFF_OPTS    = ["","FLAT","ESCALATED"]
ROOF_OPTS      = ["","Rooftop RCC","Rooftop Shed","Ground Mounted","Carport","Floating"]
INSTALL_OPTS   = ["","RCC Roof","PEB Shed","Ground Mounted","Carport","Floating Solar","Other (type manually)"]
INSTTECH_OPTS  = ["","Rooftop - RCC","Rooftop - Shed","Floating","Carport"]
EVAC_OPTS      = ["","LT","HT","11 kV","22 kV","33 kV","Other"]
NETMETER_OPTS  = ["","Net Metering","No, Behind the Meter"]
TYPEOFSYS_OPTS = ["",
    "On Grid Solar System with Net Metering",
    "On Grid Solar System with Behind the Meter Configuration",
    "Off Grid Solar System with Behind the Meter Configuration",
    "Behind the Meter Solar System with Battery Energy Storage System"]
# Auto-derived from the Net Metering input — only the first two are reachable
# automatically; the Off Grid / BESS options are there for manual override.
NETMETER_TO_SYSTYPE = {
    "Net Metering":          "On Grid Solar System with Net Metering",
    "No, Behind the Meter":  "On Grid Solar System with Behind the Meter Configuration",
    "Yes":                   "On Grid Solar System with Net Metering",
    "No":                    "On Grid Solar System with Behind the Meter Configuration",
}
SUBMITTER_OPTS = ["","Carl Aaron","Venkatesh","Harshal Jadhav","Jay Marve","Harshal Phalgaonkar","Kuladeep Thanneru"]
SUBMITTER_DATA = {
    "Carl Aaron":          "Carl Aaron | carl.aaron@fourthpartner.co | +91-9551062519",
    "Venkatesh":           "Venkatesh Panchali | venkatesh.p@fourthpartner.co | +91-9849252117",
    "Harshal Jadhav":      "Harshal Jadhav | harshal.jadhav@fourthpartner.co | +91-9766809677",
    "Jay Marve":           "Jay Marve | jay.marve@fourthpartner.co | +91-9082063404",
    "Harshal Phalgaonkar": "Harshal Phalgaonkar | harshal.p@fourthpartner.co | +91-7083647004",
    "Kuladeep Thanneru":   "Kuladeep Thanneru | kuladeep.t@fourthpartner.co | +91-8019447340",
}
DEFAULT_EXCLUSIONS = [
    "Safe roof access, construction power, water and internet by customer.",
    "Spare feeder/busbar extension by customer.",
    "All statutory and approval fees by customer.",
    "Roof strengthening and civil rectification by customer.",
    "Meter upgrade/replacement charges by customer.",
]
GST_OPTS = {
    "": "",
    "Mumbai":    "27AABCF1345M1ZL",
    "Hyderabad": "36AABCF1345M1ZK",
    "Chennai":   "33AABCF1345M1ZN",
}
ALCM_BOM = {
    "PH_085":"TopCON Bifacial N-type – ALCM (DCR)", "PH_086":"Renewsys, 585 Wp", "PH_087":"22.3%",
    "PH_088":"String Inverter", "PH_089":"Outdoor", "PH_090":"Solis, 150 kW",
    "PH_091":"Standing Seam", "PH_092":"Standing Seam; Make not specified",
    "PH_093":"Siechem", "PH_094":"Aluminium XLPE Armoured; Make not specified",
    "PH_095":"FRP / Ladder Type", "PH_096":"ACCB Panel; ABB/Schneider/Equivalent",
    "PH_097":"3 sets spare feeder extension; ABB/Schneider/Equivalent",
    "PH_098":"ACCB Panel; Energy Meter not specified",
    "PH_099":"Air Lightning Terminal; Reputed Make",
    "PH_100":"GI Earthing Electrode; Indelec / Equivalent",
    "PH_101":"Karam / Lifegear / Omkar Safety and Power Pvt. Ltd. / Equivalent",
    "PH_102":"FRP", "PH_103":"GI / Equivalent", "PH_104":"HDG", "PH_105":"HDG / Equivalent",
    "PH_106":"Wattmon / Equivalent", "PH_107":"Kipp & Zonen / Equivalent",
    "PH_108":"Surya Logics / Equivalent", "PH_109":"Manual cleaning system",
    "PH_110":"Sintex / Crompton Greaves / Kranti / Equivalent", "PH_111":"Not specified",
    "PH_112":"Not specified", "PH_113":"AC Aluminium XLPE Armoured Cable; Client Termination",
    "PH_114":"UPS; i-Ball / Equivalent",
    "PH_115":"Inverter Control Room; SCADA not specified",
    "PH_116":"Boundary and lighting; CCTV not specified",
    "PH_117":"CO₂ Fire Extinguishers, ABC Fire Extinguishers, Fire Buckets",
    "PH_118":"Zero Feeding Device; Wattmon / Equivalent", "PH_119":"Not specified",
}
ALMM_BOM = {
    "PH_085":"TopCON Bifacial N-type – ALMM (Non DCR)", "PH_086":"Renewsys, 620 Wp", "PH_087":"22.3%",
    "PH_088":"String Inverter", "PH_089":"Outdoor", "PH_090":"Solis, 150 kW",
    "PH_091":"Standing Seam", "PH_092":"Standing Seam; Make not specified",
    "PH_093":"Siechem", "PH_094":"Aluminium XLPE Armoured; Make not specified",
    "PH_095":"FRP / Ladder Type", "PH_096":"ACCB Panel; ABB/Schneider/Equivalent",
    "PH_097":"3 sets spare feeder extension; ABB/Schneider/Equivalent",
    "PH_098":"ACCB Panel; Energy Meter not specified",
    "PH_099":"Air Lightning Terminal; Reputed Make",
    "PH_100":"GI Earthing Electrode; Indelec / Equivalent",
    "PH_101":"Karam / Lifegear / Omkar Safety and Power Pvt. Ltd. / Equivalent",
    "PH_102":"FRP", "PH_103":"GI / Equivalent", "PH_104":"HDG", "PH_105":"HDG / Equivalent",
    "PH_106":"Wattmon / Equivalent", "PH_107":"Kipp & Zonen / Equivalent",
    "PH_108":"Surya Logics / Equivalent", "PH_109":"Manual cleaning system",
    "PH_110":"Sintex / Crompton Greaves / Kranti / Equivalent", "PH_111":"Not specified",
    "PH_112":"Not specified", "PH_113":"AC Aluminium XLPE Armoured Cable; Client Termination",
    "PH_114":"UPS; i-Ball / Equivalent",
    "PH_115":"Inverter Control Room; SCADA not specified",
    "PH_116":"Boundary and lighting; CCTV not specified",
    "PH_117":"CO₂ Fire Extinguishers, ABC Fire Extinguishers, Fire Buckets",
    "PH_118":"Zero Feeding Device; Wattmon / Equivalent", "PH_119":"",
}

def compute_lcoe(capacity_kwp, epc_wp_excl_gst, gst_rate=0.089, gst_credit=True,
                 project_life=25, inverter_life=5, om_cost_per_kwp=650.0,
                 om_escalation=0.05, daily_gen=3.13, degradation=0.006,
                 insurance_rate=0.0020, equip_change_rate=0.01,
                 discount_rate=0.12):
    """
    Levelised Cost of Generation (LCOG / LCOE), Rs per kWh.

        LCOE = NPV(all lifecycle costs) / NPV(all energy generated)

    Both numerator and denominator are discounted at the same rate. A simple
    (total cost / total units) ratio is NOT the levelised cost — that was the
    defect: it ignores the time value of money and O&M escalation, and
    understates LCOE materially on a 25-year asset.

    Cost stack (all in INR, nominal):
      t=0  Net capex  = capacity x 1000 x EPC/Wp x (1+GST), less GST input
                        credit if the client can claim it
      t=1..N  O&M     = O&M/kWp x capacity, escalated at om_escalation
              Insurance = insurance_rate x opening book value
                          (straight-line depreciation over project life)
              Equipment change = equip_change_rate x net capex, charged at each
                          inverter-life anniversary before end of life

    Generation: year-1 = capacity x daily_gen x 365, degrading by
    `degradation` p.a. from year 2 onward (year 1 at nameplate).

    Returns a dict so the numbers are auditable, not just a single float.
    """
    capacity_kwp = float(capacity_kwp or 0)
    if capacity_kwp <= 0 or float(daily_gen or 0) <= 0 or int(project_life or 0) <= 0:
        return {"lcoe": 0.0, "npv_cost": 0.0, "npv_gen": 0.0,
                "net_capex": 0.0, "npv_om": 0.0, "npv_ins": 0.0, "npv_repl": 0.0,
                "gen_yr1": 0.0, "total_units": 0.0}

    project_life = int(project_life)
    wp           = capacity_kwp * 1000.0
    cost_excl    = wp * float(epc_wp_excl_gst)
    gst_amount   = cost_excl * float(gst_rate)
    net_capex    = cost_excl + gst_amount - (gst_amount if gst_credit else 0.0)

    npv_cost = net_capex
    npv_om = npv_ins = npv_repl = 0.0
    npv_gen = 0.0
    total_units = 0.0
    gen_yr1 = capacity_kwp * float(daily_gen) * 365.0

    for y in range(1, project_life + 1):
        df   = (1.0 + discount_rate) ** y
        gen  = gen_yr1 * ((1.0 - degradation) ** (y - 1))
        om   = om_cost_per_kwp * capacity_kwp * ((1.0 + om_escalation) ** (y - 1))
        book = net_capex * max(0.0, 1.0 - (y - 1) / project_life)   # opening book value
        ins  = insurance_rate * book
        repl = (net_capex * equip_change_rate
                if (inverter_life and y % int(inverter_life) == 0 and y < project_life)
                else 0.0)

        npv_om   += om / df
        npv_ins  += ins / df
        npv_repl += repl / df
        npv_gen  += gen / df
        total_units += gen

    npv_cost += npv_om + npv_ins + npv_repl
    return {
        "lcoe":        round(npv_cost / npv_gen, 2) if npv_gen else 0.0,
        "npv_cost":    npv_cost,
        "npv_gen":     npv_gen,
        "net_capex":   net_capex,
        "npv_om":      npv_om,
        "npv_ins":     npv_ins,
        "npv_repl":    npv_repl,
        "gen_yr1":     gen_yr1,
        "total_units": total_units,
    }


def set_field(ph, val, pf=None):
    """
    Safe, single way to programmatically push a value into a widget-backed field.

    Never writes st.session_state["f_<PH>"] directly — that is what triggers
    Streamlit's warning:
      'The widget with key "f_PH_085" was created with a default value but also
       had its value set via the Session State API.'

    Instead: update prefilled (source of truth), stage a _sync_ value, and drop
    the stale widget key so the widget re-instantiates with the new default.
    Always follow a batch of set_field() calls with st.rerun().
    """
    store = pf if pf is not None else st.session_state.get("prefilled")
    if store is not None:
        store[ph] = val
    st.session_state[f"_sync_f_{ph}"] = val
    st.session_state.pop(f"f_{ph}", None)


def render_field(ph, label, typ, gv_fn):
    # ── Generic sync consumption (applies to EVERY field type) ──
    # A pending _sync_f_<PH> value overrides prefilled for this render only.
    _sync_key = f"_sync_f_{ph}"
    if _sync_key in st.session_state:
        cur = st.session_state.pop(_sync_key)
        st.session_state.pop(f"f_{ph}", None)   # allow value=/index= to take effect
    else:
        cur = gv_fn(ph)
    cur = "" if cur is None else str(cur)

    if typ == "DD_ROOF":
        idx = ROOF_OPTS.index(cur) if cur in ROOF_OPTS else 0
        st.selectbox(label, ROOF_OPTS, index=idx, key=f"f_{ph}")
    elif typ == "DD_INST":
        # Editable dropdown: preset options + "Other (type manually)" which
        # reveals a free-text box so any custom installation type can be entered.
        is_custom = cur not in INSTALL_OPTS and cur != ""
        idx = INSTALL_OPTS.index("Other (type manually)") if is_custom else \
              (INSTALL_OPTS.index(cur) if cur in INSTALL_OPTS else 0)
        sel = st.selectbox(label, INSTALL_OPTS, index=idx, key=f"_inst_sel_{ph}")
        if sel == "Other (type manually)":
            default_txt = cur if is_custom else ""
            typed = st.text_input(f"{label} — custom value", value=default_txt, key=f"_inst_custom_{ph}")
            st.session_state[f"f_{ph}"] = typed
        else:
            st.session_state[f"f_{ph}"] = sel
    elif typ == "DD_INSTTECH":
        idx = INSTTECH_OPTS.index(cur) if cur in INSTTECH_OPTS else 0
        st.selectbox(label, INSTTECH_OPTS, index=idx, key=f"f_{ph}")
    elif typ == "DD_EVAC":
        idx = EVAC_OPTS.index(cur) if cur in EVAC_OPTS else 0
        st.selectbox(label, EVAC_OPTS, index=idx, key=f"f_{ph}")
    elif typ == "DD_SUBMITTER":
        # Find current selection from stored full string
        cur_name = ""
        for name in SUBMITTER_DATA:
            if name in cur: cur_name = name; break
        idx = SUBMITTER_OPTS.index(cur_name) if cur_name in SUBMITTER_OPTS else 0
        val = st.selectbox(label, SUBMITTER_OPTS, index=idx, key=f"_sub_sel_{ph}")
        # Write full contact string to the PH field
        full = SUBMITTER_DATA.get(val, "")
        st.session_state[f"f_{ph}"] = full
        if full:
            st.caption(full)
    elif typ == "DD_GST":
        gst_locs = list(GST_OPTS.keys())
        cur_loc = ""
        for loc, gst in GST_OPTS.items():
            if gst and gst == cur: cur_loc = loc; break
        idx = gst_locs.index(cur_loc) if cur_loc in gst_locs else 0
        sel = st.selectbox(label, gst_locs, index=idx, key=f"_gst_sel_{ph}")
        gst_val = GST_OPTS.get(sel, "")
        st.session_state[f"f_{ph}"] = gst_val
        if gst_val: st.caption(f"GST No: {gst_val}")
    elif typ == "DD_USDINR":
        # Try to fetch live rate; fall back to stored value
        live_rate = None
        try:
            import urllib.request, json
            url = "https://open.er-api.com/v6/latest/USD"
            with urllib.request.urlopen(url, timeout=3) as resp:
                data = json.loads(resp.read())
                live_rate = round(data["rates"]["INR"], 2)
        except:
            pass
        default = str(live_rate) if live_rate else (cur or "95")
        val = st.text_input(label, value=default, key=f"f_{ph}")
        if live_rate:
            st.caption(f"Live rate fetched: ₹ {live_rate} / USD")
    elif typ == "DD_TIMELINE":
        idx = TIMELINE_OPTS.index(cur) if cur in TIMELINE_OPTS else 0
        val = st.selectbox(label, TIMELINE_OPTS, index=idx, key=f"f_{ph}")
        # Bidirectional sync: Cover timeline ↔ Project Schedule timeline.
        # CAPEX pair = PH_008 / PH_120 · OPEX pair = PH_009 / PH_100
        TIMELINE_PAIRS = {"PH_008":"PH_120","PH_120":"PH_008",
                          "PH_009":"PH_100","PH_100":"PH_009"}
        _prev_val = st.session_state.get(f"_prev_{ph}", None)
        if val != _prev_val:
            st.session_state[f"_prev_{ph}"] = val
            twin = TIMELINE_PAIRS.get(ph)
            if twin and "prefilled" in st.session_state:
                set_field(twin, val)
    elif typ == "DD_CONTRACT":
        idx = CONTRACT_OPTS.index(cur) if cur in CONTRACT_OPTS else 0
        val = st.selectbox(label, CONTRACT_OPTS, index=idx, key=f"f_{ph}")
        # OPEX: Contract Term (PH_103) drives Tariff Period (PH_105).
        # PH_105 is itself a rendered widget, so route through set_field, never
        # st.session_state["f_PH_105"] = ... (that caused the same warning).
        if ph == "PH_103":
            _tp = f"{val} Years" if val else ""
            if st.session_state.get("_prev_tariff_period") != _tp:
                st.session_state["_prev_tariff_period"] = _tp
                set_field("PH_105", _tp)
    elif typ == "DD_TARIFF":
        idx = TARIFF_OPTS.index(cur) if cur in TARIFF_OPTS else 0
        st.selectbox(label, TARIFF_OPTS, index=idx, key=f"f_{ph}")
    elif typ == "YN":
        idx = 0 if cur.lower() in ("yes","") else 1
        val = st.selectbox(label, ["Yes","No"], index=idx, key=f"f_{ph}")
        # OPEX: Net Metering (Yes/No) drives Type of System
        if ph == "PH_006":
            _mapped = NETMETER_TO_SYSTYPE.get(val, "")
            if _mapped and st.session_state.get("_prev_systype_PH_006") != val:
                st.session_state["_prev_systype_PH_006"] = val
                set_field("PH_065", _mapped)
    elif typ == "DD_NETMETER":
        idx = NETMETER_OPTS.index(cur) if cur in NETMETER_OPTS else 0
        val = st.selectbox(label, NETMETER_OPTS, index=idx, key=f"f_{ph}")
        # CAPEX: Net Metering drives Type of System
        if ph == "PH_005":
            _mapped = NETMETER_TO_SYSTYPE.get(val, "")
            if _mapped and st.session_state.get("_prev_systype_PH_005") != val:
                st.session_state["_prev_systype_PH_005"] = val
                set_field("PH_079", _mapped)
    elif typ == "DD_SYSTYPE":
        idx = TYPEOFSYS_OPTS.index(cur) if cur in TYPEOFSYS_OPTS else 0
        st.selectbox(label, TYPEOFSYS_OPTS, index=idx, key=f"f_{ph}")
    elif typ == "YN_NO":
        # Default = No
        idx = 1 if cur.lower() in ("no","") else 0
        st.selectbox(label, ["Yes","No"], index=idx, key=f"f_{ph}")
    elif typ == "DATE":
        st.text_input(label, value=date.today().strftime("%d-%m-%Y"), key=f"f_{ph}")
    elif typ == "TEXTAREA":
        st.markdown(
            f'<div style="font-size:.72rem;font-weight:800;color:var(--grey);'
            f'text-transform:uppercase;letter-spacing:.05em;margin-bottom:.4rem;">{label}</div>',
            unsafe_allow_html=True)
        st.text_area(label, value=cur, height=160, key=f"f_{ph}", label_visibility="collapsed")
    elif typ == "BULLETS":
        st.markdown(
            f'<div style="font-size:.72rem;font-weight:800;color:var(--grey);'
            f'text-transform:uppercase;letter-spacing:.05em;margin-bottom:.4rem;">{label}</div>',
            unsafe_allow_html=True)
        if f"_bullets_{ph}" not in st.session_state:
            existing = cur.strip() if cur else ""
            if existing:
                lines = [l.lstrip("•·-– ").strip() for l in existing.split("\n") if l.strip()]
            elif label.startswith("Exclusions"):
                lines = list(DEFAULT_EXCLUSIONS)
            else:
                lines = [""]
            st.session_state[f"_bullets_{ph}"] = lines
        bullets = st.session_state[f"_bullets_{ph}"]
        new_bullets = []
        for i, bline in enumerate(bullets):
            bcol1, bcol2 = st.columns([11, 1])
            with bcol1:
                v = st.text_input(f"Point {i+1}", value=bline, key=f"_b_{ph}_{i}",
                                  label_visibility="collapsed", placeholder=f"• Exclusion point {i+1}…")
                new_bullets.append(v)
            with bcol2:
                st.markdown("<div style='margin-top:1.9rem'></div>", unsafe_allow_html=True)
                if st.button("✕", key=f"_bdel_{ph}_{i}", help="Remove point"):
                    st.session_state[f"_bullets_{ph}"] = [b for j,b in enumerate(st.session_state[f"_bullets_{ph}"]) if j != i]
                    st.rerun()
        if st.button("＋ Add exclusion point", key=f"_badd_{ph}"):
            st.session_state[f"_bullets_{ph}"] = new_bullets + [""]
            st.rerun()
        st.session_state[f"f_{ph}"] = "\n".join(f"• {b}" for b in new_bullets if b.strip())
    else:
        # Pending sync values already consumed at top of render_field()
        st.text_input(label, value=cur, key=f"f_{ph}")


# ═══════════════════════════════ LANDING ═══════════════════════
if st.session_state.stage == "landing":
    st.markdown("""
    <div class="hero-wrap">
      <div class="over">Fourth Partner Energy · Solar Proposal Automation</div>
      <h1>Generate <span>Solar Proposals</span><br/>in Minutes.</h1>
      <p>Upload your BOM, enter key project numbers, and download a<br/>client-ready OPEX or CAPEX proposal document instantly.</p>
      <div class="lcards">
        <div class="lcard"><div class="n">01</div><div class="l">Upload Files</div></div>
        <div class="lcard"><div class="n">02</div><div class="l">Key Inputs</div></div>
        <div class="lcard"><div class="n">03</div><div class="l">Review Data</div></div>
        <div class="lcard"><div class="n">04</div><div class="l">Download</div></div>
      </div>
    </div>""", unsafe_allow_html=True)
    _, c, _ = st.columns([1,1,1])
    with c:
        if st.button("Get Started  →", type="primary", use_container_width=True):
            st.session_state.stage = "upload"
            st.session_state["_proposal_started_at"] = datetime.now().timestamp()
            st.rerun()
    st.markdown('<div class="foot">Fourth Partner Energy · All data stays on your computer · v1.0</div>', unsafe_allow_html=True)


# ═══════════════════════════════ UPLOAD ════════════════════════
elif st.session_state.stage == "upload":
    show_header("upload")

    # ── Project type as styled toggle ──
    st.markdown('<div class="sec">Project Type</div>', unsafe_allow_html=True)
    pt = st.radio("", ["OPEX","CAPEX"], horizontal=True, key="pt_radio",
                  label_visibility="collapsed")
    st.markdown(f"""
    <div style="display:flex;gap:.75rem;margin:.5rem 0 1.25rem;">
      <div style="flex:1;background:{'var(--orange)' if pt=='OPEX' else 'var(--grey-lt)'};
        border:2px solid {'var(--orange)' if pt=='OPEX' else 'var(--border2)'};
        border-radius:14px;padding:1.1rem 1.5rem;transition:all .2s;
        box-shadow:{'0 6px 18px var(--orange-shadow)' if pt=='OPEX' else '0 2px 6px var(--shadow)'};">
        <div style="font-family:'Nunito',sans-serif;font-size:1.15rem;font-weight:900;
          color:{'#000' if pt=='OPEX' else 'var(--text3)'} !important;">OPEX</div>
        <div style="font-size:.8rem;font-weight:800;margin-top:3px;
          color:{'#000' if pt=='OPEX' else 'var(--text3)'} !important;">
          Operating Expenditure · PPA Model</div>
      </div>
      <div style="flex:1;background:{'var(--orange)' if pt=='CAPEX' else 'var(--grey-lt)'};
        border:2px solid {'var(--orange)' if pt=='CAPEX' else 'var(--border2)'};
        border-radius:14px;padding:1.1rem 1.5rem;transition:all .2s;
        box-shadow:{'0 6px 18px var(--orange-shadow)' if pt=='CAPEX' else '0 2px 6px var(--shadow)'};">
        <div style="font-family:'Nunito',sans-serif;font-size:1.15rem;font-weight:900;
          color:{'#000' if pt=='CAPEX' else 'var(--text3)'} !important;">CAPEX</div>
        <div style="font-size:.8rem;font-weight:800;margin-top:3px;
          color:{'#000' if pt=='CAPEX' else 'var(--text3)'} !important;">
          Capital Expenditure · Ownership Model</div>
      </div>
      <div style="flex:1;background:var(--grey-lt);border:2px dashed var(--border2);
        border-radius:14px;padding:1.1rem 1.5rem;opacity:.6;cursor:not-allowed;">
        <div style="font-family:'Nunito',sans-serif;font-size:1rem;font-weight:900;
          color:var(--text3) !important;">OPEN ACCESS</div>
        <div style="font-size:.72rem;font-weight:600;margin-top:3px;color:var(--text3) !important;">
          Open Access Model<br><span style="font-size:.65rem;font-weight:700;">(Coming soon)</span></div>
      </div>
    </div>""", unsafe_allow_html=True)

    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown('<div class="sec">Upload Files</div>', unsafe_allow_html=True)
    LOCAL_WORD_TEMPLATE = LOCAL_OPEX_TEMPLATE if pt == "OPEX" else LOCAL_CAPEX_TEMPLATE
    if LOCAL_WORD_TEMPLATE:
        st.markdown(f'<div class="ibox">📝 {pt} Word template detected: <b>{os.path.basename(LOCAL_WORD_TEMPLATE)}</b></div>', unsafe_allow_html=True)

    # ── Styled upload grid ──
    col_a, col_b = st.columns(2, gap="medium")
    with col_a:
        st.markdown('<div class="upload-card-title">📊 BOM Excel <span style="color:var(--orange);font-size:.65rem;">(required)</span></div>', unsafe_allow_html=True)
        bom_f = st.file_uploader("BOM Excel", type=["xlsx"], key="bom", label_visibility="collapsed")
    with col_b:
        st.markdown('<div class="upload-card-title">📋 Proposal Input Template <span style="color:var(--text3);font-size:.65rem;">(optional — BOM alone works)</span></div>', unsafe_allow_html=True)
        pi_f = st.file_uploader("Proposal Input Template", type=["xlsx"], key="pi", label_visibility="collapsed")

    col_c, col_d = st.columns(2, gap="medium")
    with col_c:
        st.markdown('<div class="upload-card-title">🖼️ Site Image <span style="color:var(--text3);font-size:.65rem;">(optional)</span></div>', unsafe_allow_html=True)
        img_f = st.file_uploader("Site Image", type=["png","jpg","jpeg"], key="img", label_visibility="collapsed")
    with col_d:
        wt_f = None
        if not LOCAL_WORD_TEMPLATE:
            st.markdown('<div class="upload-card-title">📝 Word Template <span style="color:var(--orange);font-size:.65rem;">(required)</span></div>', unsafe_allow_html=True)
            wt_f = st.file_uploader("Word Template", type=["docx"], key="wt", label_visibility="collapsed")
        else:
            st.markdown('<div class="upload-card-title" style="color:var(--grey);">✅ Word Template</div>', unsafe_allow_html=True)
            st.markdown(f'<div style="padding:.75rem 1rem;background:var(--cyan-lt);border-radius:10px;border:1.5px solid var(--cyan-md);font-size:.8rem;font-weight:700;color:var(--grey-dk);">Auto-detected from app folder</div>', unsafe_allow_html=True)

    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    cb1, cb2 = st.columns([1,3])
    with cb1:
        if st.button("← Back", use_container_width=True): st.session_state.stage="landing"; st.rerun()
    with cb2:
        if st.button("Extract BOM & Continue →", type="primary", use_container_width=True):
            # Validate only BOM + Word template required
            miss = []
            if not bom_f: miss.append("BOM Excel")
            if not LOCAL_WORD_TEMPLATE and not wt_f: miss.append("Word Template")
            if miss: st.error(f"Missing: {', '.join(miss)}"); st.stop()

            with st.spinner("Reading BOM…"):
                import openpyxl

                st.session_state.project_type = pt
                td = tempfile.mkdtemp(); st.session_state.tmp_dir = td

                # ── Save files ──
                bp = os.path.join(td, "bom.xlsx")
                open(bp, "wb").write(bom_f.getbuffer())
                if LOCAL_WORD_TEMPLATE:
                    wp = os.path.join(td, "tpl.docx"); shutil.copy2(LOCAL_WORD_TEMPLATE, wp)
                else:
                    wp = os.path.join(td, "tpl.docx"); open(wp, "wb").write(wt_f.getbuffer())
                if img_f:
                    ip = os.path.join(td, img_f.name); open(ip, "wb").write(img_f.getbuffer())
                    st.session_state.site_image_path = ip
                else:
                    st.session_state.site_image_path = None

                st.session_state.word_template_path = wp

                # ── Fast BOM read — read_only=True, data_only=True ──
                wb_bom = openpyxl.load_workbook(bp, read_only=True, data_only=True)

                def _cell(wb, sheet_name, cell_ref):
                    """Safe cell reader — returns stripped string or empty."""
                    try:
                        ws = wb[sheet_name]
                        v = ws[cell_ref].value
                        return str(v).strip() if v is not None else ""
                    except Exception:
                        return ""

                def _num(wb, sheet_name, cell_ref):
                    """Safe numeric cell reader."""
                    try:
                        ws = wb[sheet_name]
                        v = ws[cell_ref].value
                        return float(v) if v is not None else 0.0
                    except Exception:
                        return 0.0

                BOM_SHEET    = "Bidcell BOM (LTS)"
                DESIGN_SHEET = "Input's from Initial Design"
                BD_SHEET     = "Input's from BD(Sales)"

                # ── Extract directly from BOM — corrected cell references ──
                plant_cap    = _num(wb_bom,  BOM_SHEET,    "G18")   # Plant capacity kWp
                epc_cost     = _num(wb_bom,  BOM_SHEET,    "B21")   # EPC Cost excl. GST (₹/Wp)
                module_make  = _cell(wb_bom, BOM_SHEET,    "H31")   # Module Make/Model
                dc_cable     = _cell(wb_bom, BOM_SHEET,    "H139")  # DC Cable
                module_wp    = _cell(wb_bom, DESIGN_SHEET, "F22")   # Module Wattage
                cust_name    = _cell(wb_bom, BOM_SHEET,    "G6")    # Customer Name
                address      = _cell(wb_bom, BD_SHEET,     "B13")   # Site address
                struct_type  = _cell(wb_bom, DESIGN_SHEET, "D25")   # Structure type
                inst_type    = _cell(wb_bom, DESIGN_SHEET, "C25")   # Type of Installation
                contact      = _cell(wb_bom, BD_SHEET,     "B76")   # Contact Person
                email        = _cell(wb_bom, BD_SHEET,     "B79")   # Email
                phone        = _cell(wb_bom, BD_SHEET,     "B77")   # Phone number
                mobile       = _cell(wb_bom, BD_SHEET,     "B11")   # Mobile fallback

                # Lat/Long — B14 contains "lat, long" — split on comma
                latlong_raw  = _cell(wb_bom, BD_SHEET, "B14")
                lat_val, lon_val = "", ""
                if latlong_raw and "," in latlong_raw:
                    parts = latlong_raw.split(",", 1)
                    lat_val = parts[0].strip()
                    lon_val = parts[1].strip()
                elif latlong_raw:
                    lat_val = latlong_raw.strip()

                wb_bom.close()

                # ── Build prefilled dict ──
                pf = {}

                if plant_cap:   pf["PH_001"] = plant_cap
                if epc_cost:    pf["_epc_excl_gst"] = epc_cost

                if cust_name:   pf["PH_019"] = cust_name
                if address:     pf["PH_024"] = address; pf["PH_020"] = address
                if contact:     pf["PH_021"] = contact
                if phone:       pf["PH_022"] = phone
                elif mobile:    pf["PH_022"] = mobile
                if email:       pf["PH_023"] = email
                if lat_val:     pf["PH_025"] = lat_val
                if lon_val:     pf["PH_026"] = lon_val
                if inst_type:
                    # Map to DD_INST options if possible
                    _inst_map = {"rcc":"RCC Roof","peb":"PEB Shed","ground":"Ground Mounted",
                                 "carport":"Carport","floating":"Floating Solar"}
                    _k = inst_type.lower()
                    pf["PH_027"] = next((v for k,v in _inst_map.items() if k in _k), inst_type)
                    pf["PH_040"] = pf["PH_027"]

                if module_make:
                    pf["PH_007"] = f"{module_make} / Equivalent" if "Equivalent" not in module_make else module_make
                    pf["PH_072"] = module_make
                pf.setdefault("PH_071", "TopCON Bifacial N-type")   # PV Module — Type default
                if module_wp:   pf["PH_002"] = f"{module_wp}Wp Mono PERC"
                # Inverter make → PH_007 (Key Components Inverters) and PH_076 (BOM Inverter Make)
                # PH_008 = Project Timeline (DD_TIMELINE dropdown) — NOT set here, user picks it
                pf["PH_076"] = "Solis"
                if struct_type: pf["PH_077"] = struct_type
                if dc_cable:    pf["PH_079"] = dc_cable

                # ── If PI template also uploaded, merge it (adds more fields) ──
                res, warn = [], []
                if pi_f:
                    pp = os.path.join(td, "pi.xlsx")
                    open(pp, "wb").write(pi_f.getbuffer())
                    up = os.path.join(td, "pi_updated.xlsx")
                    try:
                        ok, res, warn = transfer_values(bp, pp, up, OPEX_MAPPINGS)
                        if ok:
                            # Merge PI values — PI wins over BOM for same keys
                            wb_pi = openpyxl.load_workbook(up, read_only=True, data_only=True)
                            ws_pi = wb_pi["Proposal Inputs"]
                            for r in range(1, ws_pi.max_row + 1):
                                a = ws_pi.cell(r, 1).value
                                if isinstance(a, str) and a.strip().startswith("PH_"):
                                    v = ws_pi.cell(r, 3).value
                                    if v is not None: pf[a.strip()] = v
                            wb_pi.close()
                            st.session_state.updated_proposal_path = up
                        else:
                            for e in warn: st.warning(e)
                    except Exception as ex:
                        st.warning(f"PI template merge skipped: {ex}")

                # ── If no PI, create a minimal pi_updated from BOM data ──
                if not pi_f or not os.path.exists(os.path.join(td, "pi_updated.xlsx")):
                    # Write pf dict into a fresh minimal workbook
                    wb_out = openpyxl.Workbook()
                    ws_out = wb_out.active; ws_out.title = "Proposal Inputs"
                    for ph, val in pf.items():
                        ws_out.append([ph, "", val])
                    up = os.path.join(td, "pi_updated.xlsx"); wb_out.save(up)
                    st.session_state.updated_proposal_path = up

                st.session_state.transfer_results = res
                st.session_state.transfer_warnings = warn
                st.session_state.prefilled = pf
                st.session_state.capacity_kwp = float(pf.get("PH_001", 0) or 0)

            st.session_state.stage = "key_inputs"; st.rerun()


# ═══════════════════════════════ KEY INPUTS ════════════════════
elif st.session_state.stage == "key_inputs":
    show_header("key_inputs")
    with st.expander(f"📋 BOM Extracted — {len(st.session_state.prefilled)} fields", expanded=False):
        pf_show = st.session_state.prefilled
        bom_labels = {
            "PH_001":"Plant Capacity (kWp)","PH_019":"Customer Name",
            "PH_024":"Site Address","PH_007":"Module Make/Model",
            "PH_071":"Module Wattage","PH_008":"Inverter Make",
            "PH_076":"Inverter Model","PH_077":"Structure Type","PH_079":"DC Cable"
        }
        shown = 0
        for ph, lbl in bom_labels.items():
            if ph in pf_show and pf_show[ph]:
                st.markdown(f'<div class="cr"><span class="lb">{lbl}</span><span class="vl">{pf_show[ph]}</span></div>', unsafe_allow_html=True)
                shown += 1
        if shown == 0:
            st.info("No BOM fields extracted — check sheet names match expected format.")
        for w in st.session_state.transfer_warnings:
            st.warning(w)

    is_capex = st.session_state.get("project_type") == "CAPEX"
    _bom_cap = st.session_state.capacity_kwp   # original value from BOM
    pf_bom = st.session_state.prefilled
    # Pull EPC default from BOM if available
    _epc_default_capex = float(pf_bom.get("_epc_excl_gst", 43.95) or 43.95)
    _epc_default_opex  = float(pf_bom.get("_epc_excl_gst", 31.65) or 31.65)

    st.markdown('<div class="sec">Proposed Capacity</div>', unsafe_allow_html=True)
    c_cap1, c_cap2 = st.columns([1, 2])
    with c_cap1:
        cap = st.number_input(
            "Capacity (kWp)",
            min_value=0.01,
            value=float(_bom_cap) if _bom_cap else 0.01,
            step=0.01,
            format="%.2f",
            key="ki_cap",
            help="Pre-filled from BOM. Edit to override — this value is used for all calculations and the Word document.",
        )
    with c_cap2:
        if abs(cap - _bom_cap) > 0.001:
            st.markdown(f'<div class="ibox" style="margin-top:1.6rem">BOM value: <b>{fmt_cap(_bom_cap)} kWp</b> &nbsp;→&nbsp; Overridden to <b>{fmt_cap(cap)} kWp</b></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="cap" style="margin-top:.5rem"><div class="v">{fmt_cap(cap)} kWp</div><div class="l">From BOM — {"CAPEX" if is_capex else "OPEX"}</div></div>', unsafe_allow_html=True)
    # Update session state so the overridden cap flows everywhere
    st.session_state.capacity_kwp = cap
    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown('<div class="sec">Key Project Inputs</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        project_life = st.selectbox("Project Life (Years)", [10,15,20,25], index=3, key="ki_life")

    if not is_capex:
        with c1:
            solar_tariff = st.number_input("Our Solar Tariff (₹/kWh)",       min_value=0.00, value=3.90,                step=0.01, format="%.2f", key="ki_st")
            yield_kwh    = st.number_input("Generation Yield (kWh/kWp/day)", min_value=0.00, value=3.82,                step=0.01, format="%.2f", key="ki_y")
        with c2:
            grid_tariff  = st.number_input("Variable Grid Tariff (₹/kWh)",   min_value=0.00, value=7.50,              step=0.01, format="%.2f", key="ki_gt")
            epc_infra    = st.number_input("EPC and Infra (₹/Wp)",           min_value=0.00, value=_epc_default_opex, step=0.01, format="%.2f", key="ki_epc")
            # Financing = EPC × 70% × 1.25% — auto-calculated, user can override
            _fin_auto    = round(epc_infra * 0.70 * 0.0125, 4)
            financing    = st.number_input("Financing Cost (₹/Wp)  [auto: EPC × 70% × 1.25%]",
                                           min_value=0.00, value=_fin_auto,
                                           step=0.01, format="%.4f", key="ki_fin",
                                           help=f"Auto = {epc_infra:.2f} × 70% × 1.25% = {_fin_auto:.4f}. Edit to override.")
            inv_kw_opex  = st.selectbox("Inverter Capacity (kW)", [150, 125, 80, 60], key="ki_inv_kw_opex")
        gst              = round(epc_infra * 0.089, 2)
        epc_wp           = round(epc_infra + gst, 2)
        total_project_wp = round(epc_infra + financing, 2)   # EPC + financing (no GST on financing)
        total_cost       = total_project_wp * cap * 1000
        irradiation      = round(yield_kwh * 365 / 0.76, 0) if yield_kwh > 0 else 0
        annual_gen       = cap * yield_kwh * 365
        per_unit_sav     = grid_tariff - solar_tariff
        yr1_sav          = annual_gen * per_unit_sav
        total_sav        = yr1_sav * project_life
        inv_kw           = inv_kw_opex
    else:
        from calculations import calc_capex_financials
        with c1:
            epc_excl = st.number_input("EPC Cost excl. GST (₹/Wp)",      min_value=0.00, value=_epc_default_capex, step=0.01, format="%.2f", key="ki_epc_ex")
            daily_gen= st.number_input("Daily Generation (kWh/kWp/day)", min_value=0.00, value=3.80,               step=0.01, format="%.2f", key="ki_dg")
            om_cost  = st.number_input("O&M Cost (₹/kWp/year)",          min_value=0.0,  value=650.0,             step=10.0, format="%.0f", key="ki_om")
            inv_kw   = st.selectbox("Inverter Capacity (kW)", [150, 125, 80, 60], key="ki_inv_kw")
        with c2:
            eb_tariff= st.number_input("Current EB Tariff (₹/kWh)",      min_value=0.00, value=7.08,  step=0.01, format="%.2f", key="ki_eb")
            eb_esc   = st.number_input("EB Tariff Escalation (%)",        min_value=0.0,  value=2.0,   step=0.1,  format="%.1f", key="ki_esc")
            om_esc   = st.number_input("O&M Escalation (%)",              min_value=0.0,  value=5.0,   step=0.1,  format="%.1f", key="ki_omesc")
            inv_life = st.number_input("Inverter Life (years)",           min_value=1,    value=5,     step=1,                   key="ki_inv")
        capex_fin = calc_capex_financials(capacity_kwp=cap, epc_wp_excl_gst=epc_excl, project_life=project_life,
            inverter_life=inv_life, om_cost_per_kwp=om_cost, om_escalation=om_esc/100,
            daily_gen=daily_gen, eb_tariff=eb_tariff, eb_escalation=eb_esc/100)
        gst = round(epc_excl * 0.089, 2); epc_wp = round(epc_excl + gst, 2)
        total_cost = capex_fin["total_cost"]; annual_gen = capex_fin["gen_yr1"]
        irradiation = round(daily_gen * 365 / 0.76, 0) if daily_gen > 0 else 0
        # defaults for fields now in Estimated Generation tab
        ghi_raw = "1,964"
        degrad  = "2.5% Year 1, 0.7% p.a. thereafter"

    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown('<div class="sec">Summary</div>', unsafe_allow_html=True)
    if not is_capex:
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr-hi"><span class="lb">Capacity</span><span class="vl">{fmt_cap(cap)} kWp</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr-hi"><span class="lb">GST @8.9%</span><span class="vl">₹ {gst:.2f}</span></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="cr-hi"><span class="lb">Total EPC/Wp</span><span class="vl">₹ {epc_wp:.2f}</span></div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr"><span class="lb">Total Project Cost/Wp</span><span class="vl">₹ {total_project_wp:.2f}</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr"><span class="lb">Total Cost</span><span class="vl">₹ {fmt_indian(total_cost)}</span></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="cr"><span class="lb">Annual Gen</span><span class="vl">{fmt_indian(annual_gen)} kWh</span></div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr"><span class="lb">1st Year Savings</span><span class="vl">₹ {fmt_indian(yr1_sav)}</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr"><span class="lb">Total Savings ({project_life} yrs)</span><span class="vl">₹ {fmt_indian(total_sav)}</span></div>', unsafe_allow_html=True)
    else:
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr-hi"><span class="lb">Capacity</span><span class="vl">{fmt_cap(cap)} kWp</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr-hi"><span class="lb">Total EPC/Wp incl. GST</span><span class="vl">₹ {epc_wp:.2f}</span></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="cr"><span class="lb">Investment</span><span class="vl">₹ {capex_fin["investment_cr"]:.2f} Cr</span></div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr"><span class="lb">Generation Yr1</span><span class="vl">{fmt_indian(annual_gen)} kWh</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr"><span class="lb">Savings Yr1</span><span class="vl">₹ {capex_fin["savings_yr1_lacs"]:.2f} L</span></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="cr-hi"><span class="lb">Payback</span><span class="vl">{capex_fin["payback_years"]:.1f} yrs</span></div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1: st.markdown(f'<div class="cr"><span class="lb">Net Savings ({project_life}yr)</span><span class="vl">₹ {capex_fin["net_savings_lacs"]:.2f} L</span></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="cr"><span class="lb">Project IRR</span><span class="vl">{capex_fin["project_irr"]*100:.2f}%</span></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="cr"><span class="lb">Equity IRR</span><span class="vl">{capex_fin["equity_irr"]*100:.2f}%</span></div>', unsafe_allow_html=True)

    env = calc_environmental(annual_gen)
    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown('<div class="sec">Environmental Impact</div>', unsafe_allow_html=True)
    e1,e2,e3,e4 = st.columns(4)
    with e1: st.markdown(f'<div class="ec ec-coal"><div class="v">🪨 {fmt_indian(env["coal_tons"])} t</div><div class="l">Coal Saved / yr</div></div>', unsafe_allow_html=True)
    with e2: st.markdown(f'<div class="ec ec-co2"><div class="v">☁️ {fmt_indian(env["co2_tons"])} t</div><div class="l">CO₂ Reduced / yr</div></div>', unsafe_allow_html=True)
    with e3: st.markdown(f'<div class="ec ec-water"><div class="v">💧 {fmt_indian(int(round(env["water_litres"])))} L</div><div class="l">Water Saved / yr</div></div>', unsafe_allow_html=True)
    with e4: st.markdown(f'<div class="ec ec-trees"><div class="v">🌳 {fmt_indian(env["trees"])}</div><div class="l">Trees Equivalent</div></div>', unsafe_allow_html=True)

    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    cb1,cb2 = st.columns([1,3])
    with cb1:
        if st.button("← Back to Upload", use_container_width=True): st.session_state.stage="upload"; st.rerun()
    with cb2:
        if st.button("Review All Data & Continue →", type="primary", use_container_width=True):
            if not is_capex:
                term_vals = calc_termination(cap,epc_wp,financing,solar_tariff,0,yield_kwh,project_life)
            else:
                term_vals = []
            if is_capex:
                st.session_state.key_inputs = {"project_life":project_life,"capex_fin":capex_fin,
                    "annual_gen":annual_gen,"env":env,"term_vals":term_vals}
                pf = st.session_state.prefilled
                if not pf.get("PH_085"): pf["PH_085"] = "TopCON Bifacial N-type"   # PV Module — Type default
                if not pf.get("PH_087"): pf["PH_087"] = "22.3%"                    # PV Module — Efficiency default
                if not pf.get("PH_088"): pf["PH_088"] = "String Inverter"          # Inverter — Type default

                # ══ SECTION 1: COVER PAGE & EXECUTIVE SUMMARY ══
                # PH_001 Proposed Capacity — exact value, no rounding
                pf["PH_001"] = fmt_cap(cap)                             # exact capacity, no rounding
                pf["PH_003"] = f"{fmt_indian(int(round(annual_gen)))} kWh/year"
                _mod_make = str(pf.get("PH_072", pf.get("PH_006", "")) or "").strip()
                pf["PH_006"] = f"{_mod_make} / Equivalent" if _mod_make and "Equivalent" not in _mod_make else _mod_make
                pf["PH_007"] = "Solis / Equivalent"           # just Solis, no kW
                pf["PH_009"] = "\n".join(DEFAULT_EXCLUSIONS)
                pf["PH_017"] = date.today().strftime("%d %B %Y")
                pf["PH_018"] = "15 days from the date of proposal"
                pf["PH_030"] = "No"                          # Shading default No
                pf["PH_038"] = fmt_cap(cap)                  # Capacity Assessment — exact, no rounding
                _invest_cr = round(capex_fin.get("investment_cr", capex_fin.get("system_cost_lacs",0)/100), 2)
                pf["PH_010"] = f"₹ {_invest_cr} Cr"
                # % Consumption Replaced (PH_004) — user enters this manually on the form.
                # Seed with empty string; user fills it in Section 1 of the review form.
                pf.setdefault("PH_004", "")
                pf["PH_155"] = pf.get("PH_004", "")
                # PH_011 Payment Terms
                # NOTE: no leading "•" here — the template cell already applies
                # an auto-bullet list format per line; adding a manual bullet
                # on top produced the double-bullet issue seen in the doc.
                pf["PH_011"] = (
                    "30% Advance along with PO\n"
                    "The following payment would be through a bank guarantee/ irrevocable Letter of Credit "
                    "equivalent of 70% of the project cost\n"
                    "60% after receipt of material at site of work on pro rata basis;\n"
                    "7.5% after plant commissioning of the plant i.e. \"Commissioning Date\"\n"
                    "2.5% after receipt of CEIG approval but no later than 60 days from the \"Commissioning Date\""
                )
                # PH_012 AMC Cost
                _amc_lacs = round(capex_fin["amc_cost_lacs"], 2)
                pf["PH_012"] = f"₹ {_amc_lacs} Lakhs/year with 5% escalation"
                # PH_013 Payback Period
                pf["PH_013"] = f"{round(capex_fin['payback_years'], 2)} Years"
                # PH_014 Net Savings over 25 Years
                _net_sav_cr = round(capex_fin["net_savings_lacs"]/100, 2)
                pf["PH_014"] = f"₹ {_net_sav_cr} Cr"
                # PH_015 CO₂ Reduction
                pf["PH_015"] = fmt_indian(int(round(env["co2_tons"])))
                # PH_016/017/018 Proposal details
                pf["PH_017"] = date.today().strftime("%d %B %Y")

                # ══ SECTION 2: SITE DETAILS ══
                # PH_019–030 customer/site — already in pf from BOM extraction
                # PH_020 short address: city only
                full_addr = str(pf.get("PH_024","")).strip()
                if full_addr:
                    cleaned = re.sub(r'\b\d{6}\b', '', full_addr).strip().rstrip('.,- ')
                    parts = [p.strip() for p in re.split(r'[,]+', cleaned) if p.strip()]
                    if parts:
                        last_part = parts[-1].strip()
                        pf["PH_020"] = last_part.split()[0] if last_part else last_part

                # ══ SECTION 4: PLANT LAYOUT ══
                pf["PH_040"] = pf.get("PH_027","")   # Type of Installation → Plant Layout
                pf["PH_053"] = fmt_cap(cap)           # Rooftop total capacity — exact, no rounding
                pf["PH_064"] = 0                      # Shed total (0 unless user fills)

                # ══ SECTION 5: ESTIMATED GENERATION ══
                # Type of System — auto-derived from the Net Metering input (PH_005);
                # blank until the user picks Net Metering / Behind the Meter.
                pf["PH_079"] = NETMETER_TO_SYSTYPE.get(pf.get("PH_005",""), "")
                pf["PH_080"] = f"{fmt_indian(int(round(annual_gen)))} kWh/year"
                pf["PH_081"] = degrad                  # Degradation
                pf["PH_082"] = f"{fmt_indian(int(round(annual_gen * 0.9)))} kWh/year"
                pf["PH_083"] = project_life
                pf["PH_084"] = ghi_raw                 # Irradiation (GHI)

                # ══ SECTION 7: PROJECT SCHEDULE ══
                pf["PH_120"] = pf.get("PH_008","")    # Delivery = same as Project Timeline

                # ══ SECTION 8: PROJECT COST ══
                _epc_per_wp = round(epc_wp, 2)
                _total_invest = round(_invest_cr * 1e7, 0)  # Cr → Rs
                # EPC — Total Investment = EPC — Investment per Wp × capacity (Wp)
                _epc_total_invest = round(epc_excl * cap * 1000, 0)
                # TOTAL — Total Investment = EPC total investment + GST @ 8.9%
                _epc_total_invest_gst = round(_epc_total_invest * 1.089, 0)
                pf["PH_121"] = f"{epc_excl:.2f}"   # EPC row: raw user input (excl. GST), not the GST-inclusive total
                pf["PH_122"] = fmt_indian(int(_epc_total_invest))
                pf["PH_127"] = f"{_epc_per_wp:.2f}"
                pf["PH_128"] = fmt_indian(int(_epc_total_invest_gst))
                pf["PH_129"] = amount_to_words_indian(_epc_total_invest)  # auto-follows PH_122
                pf["PH_130"] = 95                       # USD-INR default
                pf["PH_131"] = f"{om_cost:.2f}"         # AMC ₹/kWp/yr (2 decimals)

                # ══ SECTION 9: FINANCIAL ANALYSIS ══
                pf["PH_132"] = fmt_cap(cap)        # System Size — exact, no rounding
                pf["PH_133"] = fmt_indian(round(capex_fin["system_cost_lacs"], 2))
                pf["PH_134"] = fmt_indian(round(capex_fin["gst_credit_lacs"], 2))
                pf["PH_135"] = fmt_indian(round(capex_fin["net_cost_lacs"], 2))
                _amc_lacs_val = round(capex_fin["amc_cost_lacs"], 2)
                pf["PH_136"] = f"{_amc_lacs_val:.2f}"
                pf["PH_137"] = fmt_indian(int(round(capex_fin["gen_yr1"])))
                pf["PH_138"] = capex_fin["eb_tariff"]
                # PH_139 must store a PERCENTAGE (e.g. 2.0 for "2%"), matching
                # what Section 9's "EB Tariff Escalation · % p.a." field expects
                # as its default. capex_fin["eb_escalation"] is a FRACTION
                # (0.02), since it's just an echo of the eb_escalation param
                # passed in (eb_esc/100). Storing the raw fraction here made
                # Section 9 default to ~0.02%, not 2% — i.e. Page 1's value
                # was silently NOT flowing through correctly to Section 9.
                pf["PH_139"] = capex_fin["eb_escalation"] * 100
                pf["PH_140"] = fmt_indian(round(capex_fin["savings_yr1_lacs"], 2))
                pf["PH_141"] = round(capex_fin["payback_years"], 2)
                pf["PH_142"] = project_life
                pf["PH_143"] = 5
                pf["PH_144"] = fmt_indian(round(capex_fin["net_savings_lacs"], 2))
                pf["PH_145"] = fmt_indian(round(capex_fin["total_units_lacs"], 2))
                pf["PH_146"] = round(capex_fin["project_irr"]*100, 2)
                pf["PH_147"] = round(capex_fin["equity_irr"]*100, 2)
                pf["PH_148"] = round(capex_fin["levelised_cost"], 4)

                # ══ SECTION 10: T&C ══
                # NOTE: plain number only — the doc template's own wording already
                # appends "days from the date of this proposal.", so a default of
                # "15 days" here previously produced "15 days days" in the output.
                pf["PH_149"] = "15"
                pf["PH_150"] = round(capex_fin["amc_cost_lacs"], 2)
                # NOTE: no leading "•" — the T&C Exclusions cell already applies an
                # auto-numbered (1., 2., 3. ...) list format per paragraph line.
                pf["PH_152"] = "\n".join(DEFAULT_EXCLUSIONS)

                # ══ COMPUTED ENV VARS — used across multiple sections ══
                _coal         = int(round(env["coal_tons"]))
                _co2          = int(round(env["co2_tons"]))
                _water_litres = int(round(env["water_litres"]))
                _trees        = int(round(env["trees"]))
                _water_m      = round(env["water_litres"] / 1_000_000, 1)

                # ══ SECTION 11: ENVIRONMENTAL IMPACT (table in doc body) ══
                pf["PH_153"] = fmt_indian(_co2)

                # ══ SECTION 12: COVER — PROJECT SNAPSHOT STAT BOXES ══
                # {2}  Proposed Capacity — plain number only e.g. "4,952.61"
                #      Box has no prefix/suffix. PH_001 already set above as fmt_indian(cap).
                #      No change needed — PH_001 is already the plain formatted capacity.

                # {3}  Est. Generation — plain value
                pf["PH_154"] = f"{fmt_indian(int(round(annual_gen)))} kWh/yr"

                # {4}  % Consumption Replaced — from Section 1 PH_004 (user enters on form)
                #      Already set above via pf.setdefault("PH_004","")
                #      STAT_SLOTS now points {4} directly to PH_004 — no action needed here.

                # {5}  Investment ₹ Cr — from Section 1 PH_010
                #      Box has "Rs" baked in as ts[0]; send ONLY the number+unit part.
                #      PH_010 is set as "Rs 1.24 Cr" — strip the Rs prefix before storing.
                _invest_cr = round(capex_fin.get("investment_cr",
                                   capex_fin.get("system_cost_lacs", 0) / 100), 2)
                # Store without Rs prefix — box supplies it. STAT_SLOTS reads PH_010.
                pf["PH_010"] = f"{_invest_cr:.2f} Cr"   # e.g. "1.24 Cr" — box adds Rs

                # {6}  CO2 Reduction / yr
                pf["PH_157"] = fmt_indian(_co2)

                # {7}  Project Timeline — strip trailing "months" before storing.
                #      The {7} box has " months" as a static suffix run; _box_set
                #      writes into [0..close_brace] and leaves " months" untouched.
                #      So we send "5 + 1" and the box produces "5 + 1 months".
                _tl_seed = pf.get("PH_008", "")
                pf["PH_158"] = re.sub(r"\s*months?\s*$", "", _tl_seed, flags=re.I).strip()

                # ══ SECTION 13: COVER ENV CARDS (bottom strip) ══
                # {8}  Coal — plain number, static label in template
                pf["PH_159"] = fmt_indian(_coal)
                # {9}  Water — "15.8 M"; box has static "LITRES / YEAR" label
                pf["PH_160"] = f"{_water_m} M"
                # {10} Trees — full Indian comma number
                pf["PH_161"] = fmt_indian(_trees)

                # ══ SECTION 14: ENV IMPACT PAGE 2x2 GRID ══
                # {11} Upper-Left:  Coal Conserved
                pf["PH_162"] = fmt_indian(_coal)
                # {12} Upper-Right: Trees Planted
                pf["PH_163"] = fmt_indian(_trees)
                # {13} Lower-Left:  Water Conserved — raw litres
                pf["PH_164"] = fmt_indian(_water_litres)
                # {14} Lower-Right: CO2 Avoided
                pf["PH_165"] = fmt_indian(_co2)

                st.session_state.stage="form"; st.rerun()
            st.session_state.key_inputs = {"project_life":project_life,"solar_tariff":solar_tariff,
                "grid_tariff":grid_tariff,"yield_kwh":yield_kwh,"epc_infra":epc_infra,
                "gst":gst,"epc_wp":epc_wp,"financing":financing,"total_project_wp":total_project_wp,
                "total_cost":total_cost,"irradiation":irradiation,
                "annual_gen":annual_gen,"per_unit_sav":per_unit_sav,"yr1_sav":yr1_sav,
                "total_sav":total_sav,"env":env,"term_vals":term_vals}
            pf = st.session_state.prefilled
            if not pf.get("PH_073"): pf["PH_073"] = "22.3%"          # PV Module — Efficiency default
            from calculations import num_to_words_indian
            pf["PH_017"] = date.today().strftime("%d-%m-%Y")
            pf["PH_018"] = "15 days from the date of Proposal"   # hardcoded, not shown on form
            pf["PH_074"] = "String Inverter"; pf["PH_076"] = f"Solis {inv_kw} kW"
            pf["PH_069"] = project_life
            pf["PH_012"] = solar_tariff; pf["PH_011"] = grid_tariff
            pf["PH_001"] = f"{fmt_cap(cap)} kWp"     # exact capacity, no rounding
            pf["PH_003"] = f"{fmt_indian(int(round(annual_gen)))} kWh/year"
            pf["PH_004"] = f"{fmt_indian(int(round(annual_gen * 0.9)))} kWh/year"   # guaranteed
            # PH_008 Key Components — Inverters — default Solis / Equivalent
            pf.setdefault("PH_008", "Solis / Equivalent")
            # PH_038 Proposed Capacity (Section 3) = same value entered on page 1
            pf["PH_038"] = fmt_cap(cap)
            # PH_033/035/037/039 Capacity Assessment — default notes
            pf.setdefault("PH_033", "Based on available shadow free rooftop area")
            pf.setdefault("PH_035", "Based on customer electricity consumption and expected solar generation")
            pf.setdefault("PH_037", "Based on applicable state regulatory limit")
            pf.setdefault("PH_039", "Final capacity proposed based on technical feasibility and regulatory considerations")
            # PH_101 Offtaker = Customer Name (PH_019) — auto-filled
            pf["PH_101"] = pf.get("PH_019", "")
            # PH_102 Offtaker — Registered Office = Customer Address (short) (PH_020) — auto-filled
            pf["PH_102"] = pf.get("PH_020", "")
            # PH_104 Lock-in default 5 years
            pf.setdefault("PH_104", "5")
            # PH_103 Contract Term + PH_105 Tariff Period — both from Project Life
            _pl_str = str(project_life)
            pf["PH_103"] = _pl_str
            pf["PH_105"] = f"{project_life} Years"
            set_field("PH_103", _pl_str, pf)   # push to form dropdown
            set_field("PH_105", f"{project_life} Years", pf)
            pf["PH_106"] = solar_tariff
            # PH_107 Tariff Type default Flat
            pf.setdefault("PH_107", "FLAT")
            pf["PH_109"] = grid_tariff; pf["PH_110"] = solar_tariff
            pf["PH_111"] = solar_tariff   # landed cost = solar tariff (rooftop)
            pf["PH_112"] = round(per_unit_sav, 2)
            pf["PH_113"] = fmt_indian(int(round(annual_gen)))
            pf["PH_114"] = fmt_indian(int(round(yr1_sav)))
            pf["PH_115"] = fmt_indian(int(round(total_sav)))
            pf["PH_116"] = num_to_words_indian(int(round(total_sav)))
            # PH_005 / PH_014 % Consumption Replaced — from PH_166 if provided
            _ann_cons = float(str(pf.get("PH_166", "0") or "0").replace(",", ""))
            if _ann_cons > 0:
                _pct = round(annual_gen / _ann_cons * 100, 1)
                _pct_str = f"{_pct}%"
                pf["PH_005"] = _pct_str; pf["PH_014"] = _pct_str; pf["PH_167"] = _pct_str
            else:
                pf.setdefault("PH_005", ""); pf.setdefault("PH_014", ""); pf.setdefault("PH_167", "")
            # PH_013 % Savings Yr1
            pf["PH_013"] = f"{round((grid_tariff - solar_tariff) / grid_tariff * 100, 1)}%"
            pf["PH_015"] = fmt_indian(int(round(env["co2_tons"])))
            # Section 5 Estimated Generation
            # Type of System — auto-derived from the Net Metering input (PH_006);
            # blank until the user picks Yes/No.
            pf["PH_065"] = NETMETER_TO_SYSTYPE.get(pf.get("PH_006",""), "")
            pf["PH_066"] = f"{fmt_indian(int(round(annual_gen)))} kWh/year"
            pf["PH_067"] = "2.5% Year 1, 0.7% p.a. thereafter"
            pf["PH_068"] = f"{fmt_indian(int(round(annual_gen * 0.9)))} kWh/year"
            pf["PH_070"] = fmt_indian(int(round(irradiation)))
            pf["PH_010"] = "\n".join(DEFAULT_EXCLUSIONS)
            # Environmental — cover stat boxes (exact positions from template scan)
            _coal  = int(round(env["coal_tons"]))
            _co2   = int(round(env["co2_tons"]))
            _water = int(round(env["water_litres"]))
            _trees = int(round(env["trees"]))
            _water_m = round(env["water_litres"] / 1e6, 1)
            pf["PH_118"] = fmt_indian(int(round(annual_gen)))   # Est. Gen stat box
            pf["PH_119"] = ""                                    # was Consumption% — now PH_167
            pf["PH_120"] = f"₹{fmt_indian(int(round(yr1_sav)))}"  # Savings/yr
            pf["PH_121"] = fmt_indian(_co2)                      # CO2 stat box
            pf["PH_167"] = pf.get("PH_167", "")                  # % Consumption Replaced stat box
            # Cover env cards
            pf["PH_123"] = fmt_indian(_coal)
            pf["PH_124"] = str(_water_m)                         # M litres plain number
            pf["PH_125"] = fmt_indian(_trees)
            # Env page 2×2
            pf["PH_126"] = fmt_indian(_coal)                     # UL Coal
            pf["PH_127"] = fmt_indian(_trees)                    # UR Trees
            pf["PH_128"] = fmt_indian(_water)                    # LL Water (raw litres)
            pf["PH_129"] = fmt_indian(_co2)                      # LR CO2
            pf["PH_117"] = fmt_indian(_co2)
            # Layout capacities
            pf["PH_053"] = fmt_cap(cap)
            pf["PH_064"] = fmt_cap(cap)
            pf["PH_040"] = pf.get("PH_027", ""); pf["PH_041"] = pf.get("PH_028", "")
            # Termination charges — only project_life rows populated, rest blank.
            # Final year of the project term must show 0 explicitly (contract-end
            # termination value), not a blank cell.
            for i in range(25):
                if i < project_life:
                    tv = term_vals[i] if i < len(term_vals) else 0
                    is_last_year = (i == project_life - 1)
                    if tv > 0:
                        pf[f"PH_{130+i}"] = fmt_indian(int(round(tv)))
                    elif is_last_year:
                        pf[f"PH_{130+i}"] = "0"
                    else:
                        pf[f"PH_{130+i}"] = ""
                else:
                    pf[f"PH_{130+i}"] = ""
            st.session_state.stage="form"; st.rerun()


# ═══════════════════════════════ REVIEW FORM ═══════════════════
elif st.session_state.stage == "form":
    show_header("form")
    is_capex = st.session_state.get("project_type") == "CAPEX"

    # ── Mini capacity card ──
    cap_display = st.session_state.prefilled.get("PH_001", st.session_state.get("capacity_kwp",""))
    pt_display  = "CAPEX" if is_capex else "OPEX"
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,var(--orange),#F5943A);
      border-radius:12px;padding:1rem 1.75rem;margin-bottom:1rem;
      display:flex;align-items:center;justify-content:space-between;
      box-shadow:0 4px 16px var(--orange-shadow);">
      <div>
        <div style="font-family:'Nunito',sans-serif;font-size:1.6rem;font-weight:900;
          color:#fff;letter-spacing:-.02em;">{cap_display} kWp</div>
        <div style="font-size:.72rem;color:rgba(255,255,255,.75);font-weight:700;margin-top:2px;">
          Proposed Capacity · {pt_display} Proposal</div>
      </div>
      <div style="font-size:.72rem;color:rgba(255,255,255,.65);font-weight:700;text-align:right;">
        Review & Edit<br/>before generating</div>
    </div>""", unsafe_allow_html=True)

    st.caption("Pre-filled from BOM + key inputs. Edit anything that needs correction before generating.")
    pf = st.session_state.prefilled
    def gv(ph):
        v = pf.get(ph,""); return "" if v is None else str(v)

    CAPEX_CORE = [
        "1 · Cover & Executive Summary (CAPEX)",
        "2 · Site Details",
        "3 · Capacity Assessment",
        "4 · Plant Layout",
        "5 · Estimated Generation (CAPEX)",
        "6 · Bill of Material (CAPEX)",
        "7 · Project Schedule (CAPEX)",
        "8 · Project Cost",
        "9 · Financial Analysis",
        "10 · Terms & Conditions (CAPEX)",
    ]
    OPEX_CORE = [
        "1 · Cover & Executive Summary (OPEX)",
        "2 · Site Details",
        "3 · Capacity Assessment",
        "4 · Plant Layout",
        "5 · Estimated Generation (OPEX)",
        "6 · Bill of Material (OPEX)",
        "7 · Project Schedule (OPEX)",
        "8 · Commercials (OPEX)",
        "11 · Termination Charges (OPEX)",
    ]
    CORE = CAPEX_CORE if is_capex else OPEX_CORE

    # Auto-filled fields (type "A") — shown as read-only info cards, not editable inputs
    AUTO_TYPES = {"A"}

    tabs = st.tabs(CORE)
    for tab, sn in zip(tabs, CORE):
        with tab:
            fields = SECTIONS[sn]

            # ── Section 8 CAPEX: keep EPC Total Investment + words in sync
            # with whatever the user has typed into EPC — Investment per Wp ──
            if sn == "8 · Project Cost":
                _epc_wp_raw = st.session_state.get("f_PH_121", gv("PH_121"))
                try:
                    _epc_wp_val = float(str(_epc_wp_raw).replace(",", "").strip())
                except Exception:
                    _epc_wp_val = 0.0
                _cap_val = 0.0
                try:
                    _cap_val = float(str(gv("PH_038") or gv("PH_001") or st.session_state.get("capacity_kwp", 0)).replace(",", ""))
                except Exception:
                    pass
                if _epc_wp_val > 0 and _cap_val > 0:
                    _epc_total = round(_epc_wp_val * _cap_val * 1000, 0)
                    pf["PH_122"] = fmt_indian(int(_epc_total))
                    pf["PH_129"] = amount_to_words_indian(_epc_total)
                    pf["PH_128"] = fmt_indian(int(round(_epc_total * 1.089, 0)))
                    st.session_state.prefilled = pf

            # ── Section 6: BOM preset selector buttons ──
            if sn in ("6 · Bill of Material (CAPEX)", "6 · Bill of Material (OPEX)"):
                st.markdown('<div class="sec">BOM Preset</div>', unsafe_allow_html=True)

                # Snapshot the BOM values that came from the uploaded BOM/Excel once,
                # so "Default" can genuinely restore them instead of blanking fields.
                _bom_phs = [ph for ph, _, _ in fields]
                if "_bom_original" not in st.session_state:
                    st.session_state["_bom_original"] = {p: pf.get(p, "") for p in _bom_phs}
                _bom_orig = st.session_state["_bom_original"]

                b1,b2,b3,b4 = st.columns(4)
                with b1:
                    if st.button("⚙️ Default", key="bom_default", use_container_width=True):
                        st.session_state["_bom_preset"] = "default"
                        for p in _bom_phs:
                            set_field(p, _bom_orig.get(p, ""), pf)
                        st.rerun()
                with b2:
                    if st.button("🔵 ALCM", key="bom_alcm", use_container_width=True):
                        st.session_state["_bom_preset"] = "alcm"
                        # ALCM_BOM is keyed on CAPEX PH numbers (PH_085–PH_119).
                        # The OPEX BOM uses PH_071–PH_099 for the same line items,
                        # so resolve by LABEL, not by PH number.
                        _capex_bom = SECTIONS["6 · Bill of Material (CAPEX)"]
                        _alcm_by_label = {lb: ALCM_BOM[p] for p, lb, _ in _capex_bom if p in ALCM_BOM}
                        for p, lb, _ in fields:
                            if lb in _alcm_by_label:
                                set_field(p, _alcm_by_label[lb], pf)
                        st.rerun()
                with b3:
                    if st.button("🟢 ALMM", key="bom_almm", use_container_width=True):
                        st.session_state["_bom_preset"] = "almm"
                        # ALMM_BOM is keyed on CAPEX PH numbers (PH_085–PH_119), same as
                        # ALCM_BOM — resolve by LABEL so it applies to OPEX BOM too.
                        _capex_bom = SECTIONS["6 · Bill of Material (CAPEX)"]
                        _almm_by_label = {lb: ALMM_BOM[p] for p, lb, _ in _capex_bom if p in ALMM_BOM}
                        for p, lb, _ in fields:
                            if lb in _almm_by_label:
                                set_field(p, _almm_by_label[lb], pf)
                        st.rerun()
                with b4:
                    if st.button("✏️ User Input", key="bom_user", use_container_width=True):
                        st.session_state["_bom_preset"] = "user"
                        for p in _bom_phs:
                            set_field(p, "", pf)
                        st.rerun()
                preset = st.session_state.get("_bom_preset","")
                if preset and preset != "user":
                    st.markdown(f'<div class="ibox">Active preset: <b>{preset.upper()}</b> — edit any field below to override</div>', unsafe_allow_html=True)
                st.markdown('<hr class="oline">', unsafe_allow_html=True)
            if sn == "9 · Financial Analysis":
                from calculations import calc_capex_financials
                pfl = st.session_state.prefilled
                def _num(raw, fallback):
                    """
                    Parse a prefilled value that may carry display formatting
                    ('0.0%', '1,234.56', '₹ 43.95', '25 Years', '', None).
                    Section 9 writes formatted strings back into prefilled, so the
                    defaults must survive a round trip — a bare float() raised
                    ValueError: could not convert string to float: '0.0%'.
                    """
                    try:
                        s = str(raw if raw is not None else "").strip()
                        s = re.sub(r"[^0-9.\-]", "", s)      # drop %, ₹, commas, units
                        if s in ("", "-", ".", "-."):
                            return float(fallback)
                        return float(s)
                    except Exception:
                        return float(fallback)

                _cap_fa  = _num(pfl.get("PH_132"), 0)

                # ── Defaults pulled from key inputs page ──
                # EPC Cost: read from Page 1's live widget session state directly
                # (mirrors the daily_gen pattern below). PH_121 stores the
                # GST-INCLUSIVE EPC/Wp (for the Project Cost section of the
                # document) — using it here as an "excl. GST" default caused
                # GST to be double-applied on top of an already GST-inclusive
                # figure, inflating cost and corrupting Payback/IRR/LCOE.
                _epc_raw  = st.session_state.get("ki_epc_ex", None)
                _epc_def  = float(_epc_raw) if (_epc_raw is not None and float(_epc_raw) > 0) else 43.95
                _pl_def   = int(_num(pfl.get("PH_142"), 25))
                _il_def   = int(_num(pfl.get("PH_143"), 5))
                _om_def   = _num(pfl.get("PH_131"), 650)
                # Daily gen: read from key inputs widget session state; PH_148 = levelised cost (wrong source)
                _dg_raw   = st.session_state.get("ki_dg", None)
                _dg_def   = float(_dg_raw) if (_dg_raw is not None and float(_dg_raw) > 0) else 3.13
                _eb_def   = _num(pfl.get("PH_138"), 7)
                _esc_def  = _num(pfl.get("PH_139"), 2)

                # Clamp to each widget's declared bounds so a stale prefilled
                # value can never throw StreamlitValueBelowMinError.
                _pl_def  = min(max(_pl_def, 5), 30)
                _il_def  = min(max(_il_def, 1), 15)
                _epc_def = max(_epc_def, 0.0)
                _om_def  = max(_om_def, 0.0)
                _eb_def  = max(_eb_def, 0.0)
                _esc_def = max(_esc_def, 0.0)

                st.markdown('''<style>
                .pi-hdr{background:#F47920;color:#fff;text-align:center;font-weight:900;
                  font-family:Nunito,sans-serif;font-size:.78rem;letter-spacing:.08em;
                  padding:.45rem 1rem;border-radius:6px;margin-bottom:.5rem;}
                .pi-sub{background:#E8F7FD;border:1px solid #B3E2F5;text-align:center;
                  font-weight:800;font-size:.72rem;padding:.28rem;color:#1A7FD4;margin:.3rem 0 .2rem;border-radius:4px;}
                .pi-tbl{width:100%;border-collapse:collapse;font-size:.77rem;font-family:"Nunito Sans",sans-serif;margin-bottom:.3rem;}
                .pi-tbl td{padding:4px 7px;border:1px solid #e5e7eb;vertical-align:middle;}
                .pi-tbl .lb{color:#3A3A3C;font-weight:600;width:55%;}
                .pi-tbl .un{color:#8E8E93;text-align:center;width:20%;font-size:.7rem;}
                .pi-tbl .vl{color:#2C2C2E;font-weight:800;text-align:right;width:25%;font-family:Nunito,sans-serif;}
                .pi-sec{font-size:.65rem;font-weight:900;color:#F47920;text-transform:uppercase;
                  letter-spacing:.12em;margin:.8rem 0 .25rem;border-bottom:1.5px solid #FDDFC4;
                  padding-bottom:.15rem;display:block;}
                /* compact number/select inputs in section 9 */
                div[data-testid="stNumberInput"] label,
                div[data-testid="stSelectbox"] label {
                  font-size:.68rem !important; font-weight:800 !important;
                  color:#58595B !important; text-transform:uppercase !important;
                  letter-spacing:.04em !important; margin-bottom:.1rem !important;}
                </style>''', unsafe_allow_html=True)

                fa_left, fa_right = st.columns([1, 1], gap="medium")
                with fa_left:
                    st.markdown('<div class="pi-hdr">PROJECT INPUTS &amp; ASSUMPTIONS</div>', unsafe_allow_html=True)
                    st.markdown('<div class="pi-sub">Project Data</div>', unsafe_allow_html=True)

                    # Capacity — fixed from BOM
                    st.markdown(f'''<table class="pi-tbl">
                      <tr><td class="lb">Capacity</td><td class="un">kWp</td><td class="vl">{_cap_fa:,.0f}</td></tr>
                    </table>''', unsafe_allow_html=True)

                    # Editable inputs — no yellow fill
                    fa9_epc  = st.number_input("Project cost (excls. GST)  ·  INR/Wp",  min_value=0.0, value=_epc_def,  step=0.01, format="%.2f", key="fa9_epc")
                    _gst_r   = 0.089
                    _gst_v   = round(fa9_epc * _gst_r, 5)
                    _epc_in  = round(fa9_epc + _gst_v, 5)
                    st.markdown(f'''<table class="pi-tbl">
                      <tr><td class="lb">GST</td><td class="un">INR</td><td class="vl">{_gst_v:.2f}</td></tr>
                      <tr><td class="lb">Project Cost (incls. of GST)</td><td class="un">INR/Wp</td><td class="vl">{_epc_in:.2f}</td></tr>
                    </table>''', unsafe_allow_html=True)

                    fa9_pl   = st.number_input("Project Life  ·  Years",   min_value=5, max_value=30, value=_pl_def, step=1, key="fa9_pl")
                    fa9_il   = st.number_input("Inverter Life  ·  Years",  min_value=1, max_value=15, value=_il_def, step=1, key="fa9_il")
                    fa9_om   = st.number_input("O&M Cost  ·  INR/kWp",    min_value=0.0, value=_om_def, step=10.0, format="%.0f", key="fa9_om")

                    st.markdown('''<table class="pi-tbl">
                      <tr><td class="lb">O&amp;M Escalation</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">5.0%</td></tr>
                    </table>''', unsafe_allow_html=True)

                    _dg_def  = max(0.01, _dg_def)
                    fa9_dg   = st.number_input("Daily Generation  ·  kWh/kWp/day", min_value=0.01, value=_dg_def, step=0.01, format="%.2f", key="fa9_dg")

                    st.markdown('''<table class="pi-tbl">
                      <tr><td class="lb">Insurance</td><td class="un">Book value</td><td class="vl">0.20%</td></tr>
                      <tr><td class="lb">Equipment change</td><td class="un">Project cost</td><td class="vl">1%</td></tr>
                    </table>''', unsafe_allow_html=True)

                    fa9_gst_cr = st.selectbox("GST Input Credit", ["Yes","No"], key="fa9_gstcr")

                    fa9_entity  = st.selectbox("Entity type",               ["Corporate","Institutional"], key="fa9_entity")
                    fa9_adddepr = st.selectbox("Additional depreciation (20%)", ["No","Yes"],              key="fa9_adddepr")
                    fa9_comm    = st.selectbox("Commissioning period",      ["Apr-Sept","Oct-Mar"],        key="fa9_comm")

                    st.markdown('<span class="pi-sec">Module Degradation</span>', unsafe_allow_html=True)
                    st.markdown('''<table class="pi-tbl">
                      <tr><td class="lb">Year 0</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">0%</td></tr>
                      <tr><td class="lb">Year 1</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">0.6%</td></tr>
                      <tr><td class="lb">Year 2-11</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">0.6%</td></tr>
                      <tr><td class="lb">Year 12-25</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">0.6%</td></tr>
                    </table>''', unsafe_allow_html=True)

                    st.markdown('<span class="pi-sec">AD Benefit</span>', unsafe_allow_html=True)
                    st.markdown('''<table class="pi-tbl">
                      <tr><td class="lb">New Tax Applicable</td><td class="un"></td><td class="vl">Yes</td></tr>
                      <tr><td class="lb">Depreciation Rate</td><td class="un">% p.a.</td><td class="vl" style="color:#FF0000;">40%</td></tr>
                      <tr><td class="lb">Tax Rate</td><td class="un">%</td><td class="vl" style="color:#FF0000;">25.63%</td></tr>
                    </table>''', unsafe_allow_html=True)

                    st.markdown('<span class="pi-sec">Levelised Cost Basis</span>', unsafe_allow_html=True)
                    fa9_disc = st.number_input("Discount Rate (for NPV / LCOE)  ·  %",
                                               min_value=0.0, max_value=25.0, value=12.0,
                                               step=0.5, format="%.1f", key="fa9_disc")
                    fa9_lcoe_method = st.selectbox(
                        "LCOE Method",
                        ["Discounted (NPV cost ÷ NPV units)", "Simple (total cost ÷ total units)"],
                        key="fa9_lcoe_method",
                        help=("Discounted is the technically correct levelised cost and the "
                              "basis lenders/consultants expect. Simple is undiscounted and "
                              "will read materially lower — use it only if the internal "
                              "FPEL model quotes on that basis."))

                    st.markdown('<span class="pi-sec">EB Data</span>', unsafe_allow_html=True)
                    fa9_eb  = st.number_input("Current Tariff  ·  INR/kWh",     min_value=0.0, value=_eb_def,  step=0.01, format="%.2f", key="fa9_eb")
                    fa9_esc = st.number_input("EB Tariff Escalation  ·  % p.a.", min_value=0.0, value=_esc_def, step=0.1,  format="%.1f", key="fa9_esc")

                    st.markdown('<span class="pi-sec">Leverage</span>', unsafe_allow_html=True)
                    st.markdown('''<table class="pi-tbl">
                      <tr><td class="lb">Debt</td><td class="un">%</td><td class="vl">70%</td></tr>
                      <tr><td class="lb">Interest rate</td><td class="un">%</td><td class="vl" style="background:#29ABE2;color:#fff;">10.0%</td></tr>
                      <tr><td class="lb">Tenure</td><td class="un">years</td><td class="vl">10</td></tr>
                    </table>''', unsafe_allow_html=True)

                # ── Live recalculation using edited inputs ──
                try:
                    _fa9_fin = calc_capex_financials(
                        capacity_kwp    = _cap_fa,
                        epc_wp_excl_gst = fa9_epc,
                        project_life    = fa9_pl,
                        inverter_life   = fa9_il,
                        om_cost_per_kwp = fa9_om,
                        om_escalation   = 0.05,
                        daily_gen       = fa9_dg,
                        eb_tariff       = fa9_eb,
                        eb_escalation   = fa9_esc / 100,
                    )
                    _gst_credit_lacs = _fa9_fin["gst_credit_lacs"] if fa9_gst_cr == "Yes" else 0.0
                    _net_cost_lacs   = round(_fa9_fin["system_cost_lacs"] - _gst_credit_lacs, 2)

                    # ── Levelised Cost of Generation — recomputed here on a proper
                    #    discounted-cashflow basis (see compute_lcoe docstring).
                    #    This deliberately overrides _fa9_fin["levelised_cost"].
                    _lc = compute_lcoe(
                        capacity_kwp      = _cap_fa,
                        epc_wp_excl_gst   = fa9_epc,
                        gst_rate          = _gst_r,
                        gst_credit        = (fa9_gst_cr == "Yes"),
                        project_life      = fa9_pl,
                        inverter_life     = fa9_il,
                        om_cost_per_kwp   = fa9_om,
                        om_escalation     = 0.05,
                        daily_gen         = fa9_dg,
                        degradation       = 0.006,
                        insurance_rate    = 0.0020,
                        equip_change_rate = 0.01,
                        discount_rate     = (0.0 if fa9_lcoe_method.startswith("Simple")
                                             else fa9_disc / 100.0),
                    )
                    _lc["method"] = fa9_lcoe_method
                    _lc["disc"]   = 0.0 if fa9_lcoe_method.startswith("Simple") else fa9_disc
                    _fa9_fin["levelised_cost"] = _lc["lcoe"]
                    st.session_state["_lcoe_breakdown"] = _lc

                    _fa9_vals = {
                        "sys_size":    f"{_cap_fa:,.0f}",
                        "sys_cost":    f"{_fa9_fin['system_cost_lacs']:,.2f}",
                        "gst_credit":  f"{_gst_credit_lacs:,.2f}",
                        "net_cost":    f"{_net_cost_lacs:,.2f}",
                        "amc":         f"{_fa9_fin['amc_cost_lacs']:,.2f}",
                        "gen_yr1":     f"{fmt_indian(int(round(_fa9_fin['gen_yr1'])))}",
                        "eb_tariff":   f"{fa9_eb:.2f}",
                        "eb_esc":      f"{fa9_esc:.1f}%",
                        "sav_yr1":     f"{_fa9_fin['savings_yr1_lacs']:,.2f}",
                        "payback":     f"{_fa9_fin['payback_years']:.2f}",
                        "proj_life":   f"{fa9_pl}",
                        "inv_life":    f"{fa9_il}",
                        "net_sav":     f"{_fa9_fin['net_savings_lacs']:,.2f}",
                        "tot_units":   f"{_fa9_fin['total_units_lacs']:,.2f}",
                        "proj_irr":    f"{_fa9_fin['project_irr']*100:.1f}%",
                        "eq_irr":      f"{_fa9_fin['equity_irr']*100:.1f}%",
                        "lcoe":        f"{_fa9_fin['levelised_cost']:.2f}",
                    }
                except Exception as _e:
                    # Fallback to prefilled values if calc fails
                    _fa9_vals = {
                        "sys_size":  gv("PH_132"), "sys_cost": gv("PH_133"),
                        "gst_credit":gv("PH_134"), "net_cost": gv("PH_135"),
                        "amc":       gv("PH_136"), "gen_yr1":  gv("PH_137"),
                        "eb_tariff": gv("PH_138"), "eb_esc":   gv("PH_139"),
                        "sav_yr1":   gv("PH_140"), "payback":  gv("PH_141"),
                        "proj_life": gv("PH_142"), "inv_life": gv("PH_143"),
                        "net_sav":   gv("PH_144"), "tot_units":gv("PH_145"),
                        "proj_irr":  gv("PH_146"), "eq_irr":   gv("PH_147"),
                        "lcoe":      gv("PH_148"),
                    }

                with fa_right:
                    st.markdown('<div class="sec">Financial Analysis</div>', unsafe_allow_html=True)
                    fa_rows = [
                        ("System Size",                          "kWp",           _fa9_vals["sys_size"],   False, False),
                        ("System Cost <i>(including GST)</i>",  "Rs. Lacs",      _fa9_vals["sys_cost"],   False, False),
                        ("GST Input Credit",                    "Rs. Lacs",      _fa9_vals["gst_credit"], False, False),
                        ("Net Cost to Client",                  "Rs. Lacs",      _fa9_vals["net_cost"],   True,  False),
                        ("AMC Cost",                            "Rs. Lacs/year", _fa9_vals["amc"],        False, False),
                        ("Solar Units Generated in the first year","Units / year",_fa9_vals["gen_yr1"],   False, False),
                        ("Present Power Tariff",                "Rs/ unit",      _fa9_vals["eb_tariff"],  False, False),
                        ("Avg. EB Tariff Increase",             "% p.a.",        _fa9_vals["eb_esc"],     False, False),
                        ("Savings in Year 1 <i>(post tax)</i>","Rs. Lacs",      _fa9_vals["sav_yr1"],    False, False),
                        ("Payback Period",                      "Years",         _fa9_vals["payback"],    True,  False),
                        ("Project Life",                        "Years",         _fa9_vals["proj_life"],  False, False),
                        ("Inverter Life",                       "Years",         _fa9_vals["inv_life"],   False, False),
                        ("Net Savings over Project Life",       "Rs. Lacs",      _fa9_vals["net_sav"],    False, False),
                        ("Total Units over Project Life",       "Lac Units",     _fa9_vals["tot_units"],  False, False),
                        ("Project IRR - Post Tax",              "%",             _fa9_vals["proj_irr"],   True,  False),
                        ("Equity IRR - Post Tax",               "%",             _fa9_vals["eq_irr"],     True,  False),
                        ("Levelised Cost of Generation",        "Rs / kWh",      _fa9_vals["lcoe"],       True,  True),
                    ]
                    tbl_rows = ""
                    for label, unit, val, bold, highlight in fa_rows:
                        bg  = "background:#29ABE2;color:#fff;" if highlight else ("background:#FEF3EA;" if bold else "")
                        fw  = "font-weight:700;" if bold else ""
                        clr = "color:#fff;" if highlight else ""
                        tbl_rows += f'''<tr style="{bg}">
                          <td style="padding:6px 8px;border:1px solid #ddd;font-size:.79rem;{fw}{clr}">{label}</td>
                          <td style="padding:6px 8px;border:1px solid #ddd;font-size:.73rem;text-align:center;{clr}">{unit}</td>
                          <td style="padding:6px 8px;border:1px solid #ddd;font-size:.82rem;text-align:right;font-family:Nunito,sans-serif;{fw}{clr}">{val}</td>
                        </tr>'''
                    st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-family:Nunito Sans,sans-serif;">
                      <thead><tr style="background:#F5F5F5;">
                        <th style="padding:7px 8px;border:1px solid #ddd;font-size:.72rem;font-weight:800;color:#58595B;text-align:left;text-transform:uppercase;letter-spacing:.04em;">Parameter</th>
                        <th style="padding:7px 8px;border:1px solid #ddd;font-size:.72rem;font-weight:800;color:#58595B;text-align:center;text-transform:uppercase;letter-spacing:.04em;">Unit</th>
                        <th style="padding:7px 8px;border:1px solid #ddd;font-size:.72rem;font-weight:800;color:#58595B;text-align:right;text-transform:uppercase;letter-spacing:.04em;">Value</th>
                      </tr></thead>
                      <tbody>{tbl_rows}</tbody>
                    </table>''', unsafe_allow_html=True)

                    # ── LCOE audit trail ──
                    _lcb = st.session_state.get("_lcoe_breakdown")
                    if _lcb and _lcb.get("npv_gen"):
                        with st.expander("How the Levelised Cost is derived"):
                            st.markdown(f"""
| Component | NPV @ {_lcb.get('disc', 12.0):.1f}% (₹ Lacs) |
|---|---:|
| Net capex (t=0, {'GST credit claimed' if fa9_gst_cr=='Yes' else 'GST not creditable'}) | {_lcb['net_capex']/1e5:,.2f} |
| O&M ({fa9_om:,.0f} ₹/kWp, 5% esc.) | {_lcb['npv_om']/1e5:,.2f} |
| Insurance (0.20% of book value) | {_lcb['npv_ins']/1e5:,.2f} |
| Equipment change (1% at each {fa9_il}-yr inverter life) | {_lcb['npv_repl']/1e5:,.2f} |
| **Total NPV of cost** | **{_lcb['npv_cost']/1e5:,.2f}** |

Discounted generation over {fa9_pl} years: **{_lcb['npv_gen']/1e5:,.2f} Lac units**
(undiscounted {_lcb['total_units']/1e5:,.2f} Lac units, Yr-1 {_lcb['gen_yr1']:,.0f} kWh, 0.6% p.a. degradation)

**LCOE = NPV cost ÷ NPV generation = ₹ {_lcb['lcoe']:.2f} / kWh**
""")

                # ── Persist Section 9 results into prefilled so they flow
                #    Excel → Word. Without this, the doc gets stale values
                #    and every edit made on this tab is silently discarded.
                pf["PH_132"] = fmt_cap(_cap_fa)
                pf["PH_133"] = _fa9_vals["sys_cost"]
                pf["PH_134"] = _fa9_vals["gst_credit"]
                pf["PH_135"] = _fa9_vals["net_cost"]
                pf["PH_136"] = _fa9_vals["amc"]
                pf["PH_137"] = _fa9_vals["gen_yr1"]
                # PH_138 / PH_139 are read back as defaults on every rerun, so
                # store them unformatted. Display formatting stays in _fa9_vals only.
                pf["PH_138"] = f"{fa9_eb:.2f}"
                pf["PH_139"] = f"{fa9_esc:.1f}"
                pf["PH_140"] = _fa9_vals["sav_yr1"]
                pf["PH_141"] = _fa9_vals["payback"]
                pf["PH_142"] = _fa9_vals["proj_life"]
                pf["PH_143"] = _fa9_vals["inv_life"]
                pf["PH_144"] = _fa9_vals["net_sav"]
                pf["PH_145"] = _fa9_vals["tot_units"]
                pf["PH_146"] = _fa9_vals["proj_irr"]
                pf["PH_147"] = _fa9_vals["eq_irr"]
                pf["PH_148"] = _fa9_vals["lcoe"]
                st.session_state.prefilled = pf
                st.caption("Section 9 values are saved automatically and will be written to the Word document.")

                # ── Optional: push headline numbers to the Cover / Exec Summary ──
                # Cover fields (PH_010 Investment ₹Cr, PH_012 AMC, PH_013 Payback,
                # PH_014 Net Savings) are filled from Key Inputs. If Section 9
                # inputs were edited, the cover will otherwise contradict Section 9.
                if st.button("⇧ Sync headline numbers to Cover / Exec Summary",
                             key="fa9_sync_cover"):
                    try:
                        _sys_cost_lacs = float(str(_fa9_vals["sys_cost"]).replace(",", ""))
                        set_field("PH_010", f"{_sys_cost_lacs/100:,.2f}", pf)          # ₹ Cr
                    except Exception:
                        pass
                    set_field("PH_012", _fa9_vals["amc"], pf)
                    set_field("PH_013", _fa9_vals["payback"], pf)
                    set_field("PH_014", _fa9_vals["net_sav"], pf)
                    st.session_state.prefilled = pf
                    st.success("Cover figures updated from Section 9. Verify units on the Cover tab (₹ Cr vs ₹ Lacs) before generating.")
                continue  # skip default field rendering for section 9

            # ── Section 7: Delivery & Installation — show Gantt chart below ──
            if sn == "7 · Project Schedule (CAPEX)":
                fields_7 = SECTIONS[sn]
                auto_7   = [(ph,lb,tp) for ph,lb,tp in fields_7 if tp == "A"]
                manual_7 = [(ph,lb,tp) for ph,lb,tp in fields_7 if tp != "A"]
                if auto_7:
                    st.markdown('<div class="sec">Auto-calculated</div>', unsafe_allow_html=True)
                    ac1, ac2 = st.columns(2)
                    for i,(ph,lb,_) in enumerate(auto_7):
                        with (ac1 if i%2==0 else ac2):
                            st.markdown(f'<div class="cr"><span class="lb">{lb}</span><span class="vl">{gv(ph) or "—"}</span></div>', unsafe_allow_html=True)
                            if f"f_{ph}" not in st.session_state:
                                st.session_state[f"f_{ph}"] = gv(ph)
                if manual_7:
                    cols7 = st.columns(2); col_idx = 0
                    for ph, label, typ in manual_7:
                        with cols7[col_idx % 2]: render_field(ph, label, typ, gv)
                        col_idx += 1
                # ── Gantt chart — pure HTML, no soffice dependency ──
                _timeline_val = st.session_state.get("f_PH_120") or st.session_state.get("f_PH_008") or gv("PH_120") or gv("PH_008") or ""

                # Gantt bar positions read from Excel (col index: C=0, D=1, E=2, F=3, G=4, H=5)
                # Row 4-14 = task rows 1-11; filled cells = orange bar
                # Bar positions read directly (cell-fill scan) from the 4 reference
                # sheets in Solar_Project_Gantt.xlsx.xlsx — "Up to 100 kWp (3+1)",
                # "100-1000 kWp (4+1)", "Above 1000 kWp (5+1)", "Above 1000 kWp (6+1)".
                _GANTT_DATA = {
                    "3 + 1 Month": {
                        "n_months": 4,
                        "bars": {
                            1:[0],2:[0],3:[0],4:[0],5:[0],
                            6:[0,1],7:[1],8:[1,2],
                            9:[2,3],10:[3],11:[3],
                        },
                        "footer_end_col": 3,
                    },
                    "4 + 1 Month": {
                        "n_months": 5,
                        "bars": {
                            1:[0],2:[0],3:[0],4:[0],5:[0,1],
                            6:[1],7:[1],8:[2,3],
                            9:[3,4],10:[4],11:[4],
                        },
                        "footer_end_col": 4,
                    },
                    "5 + 1 Month": {
                        "n_months": 6,
                        "bars": {
                            1:[0],2:[0],3:[0],4:[0],5:[0,1],
                            6:[1],7:[1],8:[2,3,4],
                            9:[4,5],10:[5],11:[5],
                        },
                        "footer_end_col": 5,
                    },
                    "6 + 1 Month": {
                        "n_months": 7,
                        "bars": {
                            1:[0],2:[0],3:[0],4:[0],5:[0,1],
                            6:[1],7:[1,2],8:[2,3,4],
                            9:[4,5,6],10:[4,5,6],11:[5,6],
                        },
                        "footer_end_col": 6,
                    },
                }
                _TASKS = [
                    "Receipt of PO and Advance Order",
                    "Customer Kick-Off Meeting",
                    "Design Approval / Site Handover (whichever later)",
                    "Submission & Approval of Design",
                    "Placement of Orders on Vendors (Material Wise)",
                    "Manufacturing & Dispatch of Materials",
                    "Mobilization of Team at Site",
                    "Installation & Test Charge",
                    "Approvals",
                    "Commercial Operation Date (COD)",
                    "Site and Documentation Handover",
                ]

                if _timeline_val in _GANTT_DATA:
                    from datetime import date as _date
                    from dateutil.relativedelta import relativedelta

                    st.markdown('<hr class="oline">', unsafe_allow_html=True)
                    st.markdown('<div class="sec">Project Timeline — Gantt</div>', unsafe_allow_html=True)

                    _gd        = _GANTT_DATA[_timeline_val]
                    _nm        = _gd["n_months"]
                    _bars      = _gd["bars"]
                    _today     = _date.today()
                    _months    = [(_today + relativedelta(months=i+1)) for i in range(_nm)]
                    _mo_labels = [m.strftime("%b-%y") for m in _months]
                    _start_lbl = _months[0].strftime("%b-%y")
                    _end_lbl   = _months[-1].strftime("%b-%y")

                    _ORANGE = "#E97132"
                    _NAVY   = "#1F3864"
                    _LGRAY  = "#F2F2F2"
                    _WHITE  = "#FFFFFF"

                    # Column widths — match Image 3 proportions exactly
                    # Sr=4%, Task=34%, remaining split equally across month cols
                    _sr_pct   = "4%"
                    _task_pct = "34%"
                    _mo_pct   = f"{int(62/_nm)}%"

                    _th  = f"padding:5px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.72rem;text-align:center;border:1px solid #fff;"
                    _thl = f"padding:5px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.72rem;text-align:left;border:1px solid #fff;"
                    _td_num   = "padding:3px 6px;font-size:.70rem;text-align:center;font-weight:700;border:1px solid #D9D9D9;"
                    _td_task  = "padding:3px 6px;font-size:.70rem;border:1px solid #D9D9D9;"
                    _td_on    = f"border:1px solid #D9D9D9;background:{_ORANGE};"
                    _td_off   = f"border:1px solid #D9D9D9;background:{_WHITE};"
                    _td_foot  = f"padding:4px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.70rem;border:1px solid #fff;"

                    # Col group for widths
                    _colgroup = f'<colgroup><col style="width:{_sr_pct}"><col style="width:{_task_pct}">'
                    for _ in range(_nm):
                        _colgroup += f'<col style="width:{_mo_pct}">'
                    _colgroup += '</colgroup>'

                    # Header rows — no "Today" row; just column headers matching Image 3
                    _col_header = f'<tr><th style="{_th}">Sr.</th><th style="{_thl}">Item Description</th>'
                    for lbl in _mo_labels:
                        _col_header += f'<th style="{_th}">{lbl}</th>'
                    _col_header += '</tr>'

                    # Task rows
                    _task_rows = ""
                    for i, task in enumerate(_TASKS):
                        _bg = _LGRAY if i % 2 == 0 else _WHITE
                        _filled = _bars.get(i + 1, [])
                        _row = f'<tr><td style="{_td_num}background:{_bg};">{i+1}</td>'
                        _row += f'<td style="{_td_task}background:{_bg};">{task}</td>'
                        for col_i in range(_nm):
                            _row += f'<td style="{_td_on if col_i in _filled else (_td_off if i%2==0 else f"border:1px solid #D9D9D9;background:{_WHITE};")}"></td>'
                        _row += '</tr>'
                        _task_rows += _row

                    # Footer — Image 3 style: left="Start Date = Sep-26", middle=empty navy, right="Project End Date = Jan-27"
                    _foot = f'<tr><td colspan="2" style="{_td_foot}text-align:right;">Start Date =</td>'
                    for col_i in range(_nm):
                        if col_i == 0:
                            _foot += f'<td style="{_td_foot}text-align:center;">{_start_lbl}</td>'
                        elif col_i == _nm - 1:
                            _foot += f'<td style="{_td_foot}text-align:center;">{_end_lbl}</td>'
                        else:
                            _foot += f'<td style="{_td_foot}text-align:right;">{"Project End Date =" if col_i == _nm-2 else ""}</td>'
                    _foot += '</tr>'

                    _gantt_html = f'''<div style="overflow-x:auto;margin-top:.5rem;">
                      <table style="width:900px;border-collapse:collapse;font-family:\'Nunito Sans\',sans-serif;table-layout:fixed;">
                        {_colgroup}
                        <thead>{_col_header}</thead>
                        <tbody>{_task_rows}</tbody>
                        <tfoot>{_foot}</tfoot>
                      </table>
                    </div>'''
                    st.markdown(_gantt_html, unsafe_allow_html=True)
                continue

            # ── OPEX Project Schedule — Gantt ──
            if sn == "7 · Project Schedule (OPEX)":
                fields_7o = SECTIONS[sn]
                manual_7o = [(ph,lb,tp) for ph,lb,tp in fields_7o if tp != "A"]
                if manual_7o:
                    cols7o = st.columns(2); col_idx = 0
                    for ph, label, typ in manual_7o:
                        with cols7o[col_idx % 2]: render_field(ph, label, typ, gv)
                        col_idx += 1

                _timeline_val_o = (st.session_state.get("f_PH_100") or gv("PH_100") or "").strip()

                # OPEX Gantt data — exact from Excel (Solar_Project_Gantt_-_OPEX_.xlsx)
                # Tasks differ per sheet; stored per timeline option
                _OPEX_GANTT = {
                    "4 + 1 Month": {
                        "n_months": 4,
                        "tasks": [
                            "PPA Signing",
                            "Customer Kick-Off Meeting",
                            "Site Handover / Design Approval (whichever later)",
                            "Mobilization of team at site",
                            "Procurement",
                            "Civil Work, Installation & Cabling",
                            "Pre Commissioning & Test Charge",
                            "Liasoning",
                            "Site Acceptance Test",
                            "Handover",
                            "Project Closure",
                        ],
                        "bars": {
                            1:[0], 2:[0], 3:[0], 4:[0],
                            5:[0,1], 6:[1,2], 7:[2],
                            8:[2,3], 9:[2,3], 10:[2,3], 11:[],
                        },
                    },
                    "5 + 1 Month": {
                        "n_months": 5,
                        "tasks": [
                            "PPA Signing",
                            "Customer Kick-Off Meeting",
                            "Site Handover / Design Approval (whichever later)",
                            "Procurement",
                            "Mobilization of team at site",
                            "Civil Work, Installation & Cabling",
                            "Pre Commissioning & Test Charge",
                            "Liasoning",
                            "Site Acceptance Test",
                            "Handover",
                            "Project Closure",
                        ],
                        "bars": {
                            1:[0], 2:[0], 3:[0], 4:[0], 5:[0],
                            6:[1,2,3], 7:[2,3], 8:[2,3],
                            9:[3,4], 10:[4], 11:[],
                        },
                    },
                    "6 + 1 Month": {
                        "n_months": 6,
                        "tasks": [
                            "PPA Signing",
                            "Customer Kick-Off Meeting",
                            "Site Handover / Design Approval (whichever later)",
                            "Mobilization of team at site",
                            "Procurement",
                            "Civil Work, Installation & Cabling",
                            "Pre Commissioning & Test Charge",
                            "Liasoning",
                            "Site Acceptance Test",
                            "Handover",
                            "Project Closure",
                        ],
                        "bars": {
                            1:[0], 2:[0], 3:[0], 4:[0],
                            5:[0,1], 6:[1], 7:[1],
                            8:[2,3,4], 9:[4,5], 10:[5], 11:[],
                        },
                    },
                }

                if _timeline_val_o in _OPEX_GANTT:
                    from datetime import date as _date
                    from dateutil.relativedelta import relativedelta
                    st.markdown('<hr class="oline">', unsafe_allow_html=True)
                    st.markdown('<div class="sec">Project Timeline — Gantt</div>', unsafe_allow_html=True)

                    _ogd       = _OPEX_GANTT[_timeline_val_o]
                    _onm       = _ogd["n_months"]
                    _obars     = _ogd["bars"]
                    _otasks    = _ogd["tasks"]
                    _otoday    = _date.today()
                    _omonths   = [(_otoday + relativedelta(months=i+1)) for i in range(_onm)]
                    _omo_lbs   = [m.strftime("%b-%y") for m in _omonths]
                    _ostart    = _omonths[0].strftime("%b-%y")
                    _oend      = _omonths[-1].strftime("%b-%y")

                    _ORANGE = "#E97132"; _NAVY = "#1F3864"; _LGRAY = "#F2F2F2"; _WHITE = "#FFFFFF"
                    _th  = f"padding:5px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.72rem;text-align:center;border:1px solid #fff;"
                    _thl = f"padding:5px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.72rem;text-align:left;border:1px solid #fff;"
                    _td_num  = "padding:3px 6px;font-size:.70rem;text-align:center;font-weight:700;border:1px solid #D9D9D9;"
                    _td_task = "padding:3px 6px;font-size:.70rem;border:1px solid #D9D9D9;"
                    _td_on   = f"border:1px solid #D9D9D9;background:{_ORANGE};"
                    _td_foot = f"padding:4px 6px;background:{_NAVY};color:#fff;font-weight:700;font-size:.70rem;border:1px solid #fff;"

                    _sr_pct = "4%"; _task_pct = "34%"; _mo_pct = f"{int(62/_onm)}%"
                    _cg = f'<colgroup><col style="width:{_sr_pct}"><col style="width:{_task_pct}">'
                    for _ in range(_onm): _cg += f'<col style="width:{_mo_pct}">'
                    _cg += '</colgroup>'

                    _hdr = f'<tr><th style="{_th}">Sr.</th><th style="{_thl}">Item Description</th>'
                    for lbl in _omo_lbs: _hdr += f'<th style="{_th}">{lbl}</th>'
                    _hdr += '</tr>'

                    _rows = ""
                    for i, task in enumerate(_otasks):
                        _bg = _LGRAY if i%2==0 else _WHITE
                        _row = f'<tr><td style="{_td_num}background:{_bg};">{i+1}</td>'
                        _row += f'<td style="{_td_task}background:{_bg};">{task}</td>'
                        for ci in range(_onm):
                            _filled = ci in _obars.get(i+1, [])
                            _row += f'<td style="{_td_on if _filled else f"border:1px solid #D9D9D9;background:{_bg};"}"></td>'
                        _row += '</tr>'
                        _rows += _row

                    _foot = f'<tr><td colspan="2" style="{_td_foot}text-align:right;">Start Date =</td>'
                    for ci in range(_onm):
                        if ci == 0:
                            _foot += f'<td style="{_td_foot}text-align:center;">{_ostart}</td>'
                        elif ci == _onm-2 and _onm > 2:
                            _foot += f'<td style="{_td_foot}text-align:right;">Project End Date =</td>'
                        elif ci == _onm-1:
                            _foot += f'<td style="{_td_foot}text-align:center;">{_oend}</td>'
                        else:
                            _foot += f'<td style="{_td_foot}"></td>'
                    _foot += '</tr>'

                    st.markdown(f'''<div style="overflow-x:auto;margin-top:.5rem;">
                      <table style="width:900px;border-collapse:collapse;font-family:\'Nunito Sans\',sans-serif;table-layout:fixed;">
                        {_cg}<thead>{_hdr}</thead><tbody>{_rows}</tbody><tfoot>{_foot}</tfoot>
                      </table></div>''', unsafe_allow_html=True)
                elif not _timeline_val_o:
                    st.info("Select a Project Timeline above to display the Gantt chart.")
                continue
            if sn == "2 · Site Details":
                cols = st.columns(2); col_idx = 0
                for ph, label, typ in fields:
                    with cols[col_idx % 2]: render_field(ph, label, typ, gv)
                    col_idx += 1
                st.markdown('<hr class="oline">', unsafe_allow_html=True)
                if st.button("💾 Save & Sync to Plant Layout", key="sync_site", type="primary"):
                    pf["PH_027"] = st.session_state.get("f_PH_027", gv("PH_027"))
                    pf["PH_040"] = pf["PH_027"]
                    st.session_state.prefilled = pf
                    st.success("✓ Type of Installation synced to Plant Layout.")
                continue

            # ── Termination Charges: render as a 2-column table ──
            if sn == "11 · Termination Charges (OPEX)":
                _ki = st.session_state.get("key_inputs", {})
                _pl = _ki.get("project_life", 25)
                st.markdown(f'<div class="sec">Termination Charges — {_pl} Year Project</div>', unsafe_allow_html=True)
                _term_rows_html = ""
                for _y in range(1, _pl + 1):
                    _tv = gv(f"PH_{129 + _y}")
                    _term_rows_html += f"<tr><td style='padding:5px 10px;border:1px solid #e5e7eb;font-weight:600'>Year {_y}</td><td style='padding:5px 10px;border:1px solid #e5e7eb;text-align:right;font-family:Nunito,sans-serif;font-weight:700'>₹ {_tv} /kWp</td></tr>"
                st.markdown(f"""<div style='max-width:420px;margin:0 auto;'>
                <table style='width:100%;border-collapse:collapse;font-size:.82rem;font-family:"Nunito Sans",sans-serif;'>
                    <thead><tr>
                      <th style='padding:6px 10px;background:#F47920;color:#fff;text-align:left;border-radius:4px 0 0 0'>Year</th>
                      <th style='padding:6px 10px;background:#F47920;color:#fff;text-align:right;border-radius:0 4px 0 0'>Termination Value (₹/kWp)</th>
                    </tr></thead>
                    <tbody>{_term_rows_html}</tbody>
                </table></div>""", unsafe_allow_html=True)
                st.caption("Values auto-calculated from Key Inputs. Final year = ₹0.")
                continue

            # ── Section 1 OPEX: PH_166 annual consumption + PH_167 auto-calc ──
            if sn == "1 · Cover & Executive Summary (OPEX)":
                # PH_101 Offtaker = Customer Name — auto-sync on every render
                _cust_name = st.session_state.get("f_PH_019") or pf.get("PH_019", "")
                if _cust_name and pf.get("PH_101") != _cust_name:
                    pf["PH_101"] = _cust_name
                    st.session_state.prefilled = pf

            # ── Section 2 OPEX: PH_102 Offtaker Registered Office = Customer Address (short) ──
            if sn == "2 · Site Details":
                _cust_addr = st.session_state.get("f_PH_020") or pf.get("PH_020", "")
                if _cust_addr and pf.get("PH_102") != _cust_addr:
                    pf["PH_102"] = _cust_addr
                    st.session_state.prefilled = pf

            # ── Section 8 OPEX (Commercials): PH_116 Total Savings in words ──
            if sn == "8 · Commercials (OPEX)":
                _ts_raw = pf.get("PH_115", "")
                if _ts_raw:
                    from calculations import num_to_words_indian as _ntw
                    try:
                        _ts_num = int(str(_ts_raw).replace(",", ""))
                        pf["PH_116"] = _ntw(_ts_num)
                        st.session_state.prefilled = pf
                    except Exception:
                        pass

            # ── Default: render all fields, auto fields as read-only cards ──
            auto_fields   = [(ph,lb,tp) for ph,lb,tp in fields if tp == "A"]
            manual_fields = [(ph,lb,tp) for ph,lb,tp in fields if tp != "A"]

            # For Cover tab: split manual fields at PH_016 to add divider before submission fields
            SUBMISSION_PHS = {"PH_016","PH_017"} if not is_capex else {"PH_016","PH_017","PH_018"}
            # PH_167 is handled inline after PH_166 — exclude from normal rendering
            SKIP_PHS = {"PH_167"} if not is_capex else set()
            if sn in ("1 · Cover & Executive Summary (CAPEX)", "1 · Cover & Executive Summary (OPEX)"):
                manual_pre  = [(ph,lb,tp) for ph,lb,tp in manual_fields if ph not in SUBMISSION_PHS and ph not in SKIP_PHS]
                manual_sub  = [(ph,lb,tp) for ph,lb,tp in manual_fields if ph in SUBMISSION_PHS]
            else:
                manual_pre  = manual_fields
                manual_sub  = []

            if auto_fields:
                st.markdown('<div class="sec">Auto-calculated</div>', unsafe_allow_html=True)
                ac1, ac2 = st.columns(2)
                for i, (ph, lb, _) in enumerate(auto_fields):
                    with (ac1 if i%2==0 else ac2):
                        st.markdown(f'<div class="cr"><span class="lb">{lb}</span><span class="vl">{gv(ph) or "—"}</span></div>', unsafe_allow_html=True)
                        if f"f_{ph}" not in st.session_state:
                            st.session_state[f"f_{ph}"] = gv(ph)

            if manual_pre:
                if auto_fields: st.markdown('<div class="sec">Review & Edit</div>', unsafe_allow_html=True)
                cols = st.columns(2); col_idx = 0
                for ph, label, typ in manual_pre:
                    with cols[col_idx % 2]: render_field(ph, label, typ, gv)
                    col_idx += 1
                    # After PH_166 (Annual Consumption), show PH_167 auto-calc
                    if ph == "PH_166" and not is_capex:
                        _cons_raw = st.session_state.get("f_PH_166") or gv("PH_166") or ""
                        _cons_val = 0.0
                        try: _cons_val = float(str(_cons_raw).replace(",",""))
                        except Exception: pass
                        _ki2 = st.session_state.get("key_inputs", {})
                        _ag2 = _ki2.get("annual_gen", 0)
                        if _cons_val > 0 and _ag2 > 0:
                            _pct2 = round(_ag2 / _cons_val * 100, 1)
                            _pct2_str = f"{_pct2}%"
                            pf["PH_167"] = _pct2_str
                            pf["PH_005"] = _pct2_str
                            pf["PH_014"] = _pct2_str
                            st.session_state.prefilled = pf
                            with cols[col_idx % 2]:
                                st.markdown(f'<div class="cr-hi"><span class="lb">% Consumption Replaced (auto)</span><span class="vl">{_pct2_str}</span></div>', unsafe_allow_html=True)
                            col_idx += 1
                        else:
                            pf.setdefault("PH_167", ""); pf.setdefault("PH_005", "")
                            with cols[col_idx % 2]:
                                st.markdown('<div class="cr"><span class="lb">% Consumption Replaced (auto)</span><span class="vl">— enter consumption above</span></div>', unsafe_allow_html=True)
                            col_idx += 1

            if manual_sub:
                st.markdown('<hr class="oline">', unsafe_allow_html=True)
                st.markdown('<div class="sec">Proposal Submission</div>', unsafe_allow_html=True)
                cols = st.columns(2); col_idx = 0
                for ph, label, typ in manual_sub:
                    with cols[col_idx % 2]: render_field(ph, label, typ, gv)
                    col_idx += 1

    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    cb1,cb2 = st.columns([1,3])
    with cb1:
        if st.button("← Back to Key Inputs", use_container_width=True): st.session_state.stage="key_inputs"; st.rerun()
    with cb2:
        if st.button("🚀  Generate Proposal", type="primary", use_container_width=True):
            import openpyxl
            wb3 = openpyxl.load_workbook(st.session_state.updated_proposal_path)
            ws3 = wb3["Proposal Inputs"]
            ph_rows = {}
            for r in range(1, ws3.max_row+1):
                a = ws3.cell(r,1).value
                if isinstance(a,str) and a.strip().startswith("PH_"): ph_rows[a.strip()]=r

            # Collect all values: session state widgets + prefilled dict
            all_vals = dict(st.session_state.prefilled)  # start with prefilled
            for ph in ALL_PH:
                widget_val = st.session_state.get(f"f_{ph}")
                if widget_val is not None and str(widget_val).strip():
                    all_vals[ph] = widget_val
            # CAPEX extra keys
            for extra in ["PH_012_amc","PH_012_payment"]:
                v = st.session_state.prefilled.get(extra) or st.session_state.get(f"f_{extra}")
                if v: all_vals[extra] = v

            # Write all values to Excel
            for ph, val in all_vals.items():
                if val is None or str(val).strip() == "": continue
                if ph in ph_rows:
                    cell = ws3.cell(ph_rows[ph], 3)
                    sval = str(val)
                    # Only convert to float if it's a pure number (no ₹, letters, spaces beyond digits)
                    clean = sval.replace(",","").strip()
                    try:
                        if re.match(r'^-?\d+\.?\d*$', clean):
                            cell.value = float(clean)
                        else:
                            cell.value = sval
                    except:
                        cell.value = sval
                else:
                    ws3.append([ph, "", str(val)])

            if st.session_state.site_image_path and "PH_031" in ph_rows:
                ws3.cell(ph_rows["PH_031"],3).value = str(st.session_state.site_image_path)
            fp = os.path.join(st.session_state.tmp_dir,"pi_final.xlsx"); wb3.save(fp)
            st.session_state.final_input_path=fp; st.session_state.stage="generate"; st.rerun()


# ═══════════════════════════════ GENERATE ══════════════════════
elif st.session_state.stage == "generate":
    show_header("generate")
    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown("### Generating your proposal...")
    prog = st.progress(0, text="Preparing...")
    try:
        is_capex = st.session_state.get("project_type") == "CAPEX"
        script_name = "update_capex_proposal.py" if is_capex else "update_opex_proposal.py"
        import importlib.util
        spec = importlib.util.spec_from_file_location("gen", os.path.join(APP_DIR,script_name))
        gen  = importlib.util.module_from_spec(spec)
        gen.BASE = st.session_state.tmp_dir; gen.EXCEL = st.session_state.final_input_path
        spec.loader.exec_module(gen)
        gen.BASE = st.session_state.tmp_dir; gen.EXCEL = st.session_state.final_input_path
        prog.progress(10, text="[1/6] Reading Excel inputs...")
        d = gen.read_excel(st.session_state.final_input_path)
        cust  = gen.g(d,"PH_019","Customer"); cap_s = gen.g(d,"PH_001","XXX")

        # ── Resolve deferred values from form widget session state ──
        # PH_158 = Project Timeline on the cover snapshot card.
        # The DD_TIMELINE widget (f_PH_008 or f_PH_120) is only set on the form
        # page, AFTER key-inputs. Pull the live widget value now.
        _tl_live = (
            st.session_state.get("f_PH_120") or
            st.session_state.get("f_PH_008") or
            st.session_state.prefilled.get("PH_120","") or
            st.session_state.prefilled.get("PH_008","")
        ).strip()
        if _tl_live:
            # Strip trailing "months" — the {7} box has " months" as static suffix
            _tl_stored = re.sub(r"\s*months?\s*$", "", _tl_live, flags=re.I).strip()
            d["PH_158"] = _tl_stored
            st.session_state.prefilled["PH_158"] = _tl_stored
        # PH_155 = % Consumption Replaced on snapshot card — sync from PH_004 form widget
        _pct_live = (
            st.session_state.get("f_PH_004") or
            st.session_state.prefilled.get("PH_004","")
        ).strip()
        if _pct_live:
            d["PH_155"] = _pct_live
            st.session_state.prefilled["PH_155"] = _pct_live
        prog.progress(25, text="[2/6] Loading Word template...")
        from docx import Document; doc = Document(st.session_state.word_template_path)
        prog.progress(40, text="[3/6] Filling tables & snapshots..."); gen.fill_document(doc,d)

        # ── BOM "Qty (max)" column — only when the ALCM/ALMM preset was used ──
        # Solar PV Module / Inverter / Module Mounting Structure -> "Equiv. to
        # <capacity> kWp"; every other row that still shows the template's "xx…"
        # placeholder -> "As Per Design".
        if st.session_state.get("_bom_preset") in ("alcm", "almm"):
            _EQUIV_ROWS = {"Solar PV Module", "Inverter", "Module Mounting Structure"}
            for _t in doc.tables:
                _hdr = [c.text.strip() for c in _t.rows[0].cells]
                if _hdr[:2] != ["Sr.", "Equipment"] or "Qty (max)" not in _hdr:
                    continue
                _qty_col = _hdr.index("Qty (max)")
                _last_equip = ""
                for _row in _t.rows[1:]:
                    _cells_txt = [c.text.strip() for c in _row.cells]
                    if len(set(_cells_txt)) == 1:
                        continue   # merged section-banner row (all cells identical)
                    _equip = _cells_txt[1] or _last_equip
                    _last_equip = _equip or _last_equip
                    _qcell = _row.cells[_qty_col]
                    if _equip in _EQUIV_ROWS:
                        gen.set_cell_text(_qcell, f"Equiv. to {cap_s} kWp")
                    elif _qcell.text.strip().lower().startswith("xx"):
                        gen.set_cell_text(_qcell, "As Per Design")

        prog.progress(55, text="[4/6] Swapping text...")
        gen.do_text_swaps(doc,d)
        gen.fill_cover_stats(doc,d,verbose=False)
        if not is_capex: gen.fill_termination_charges(doc,d)
        prog.progress(70, text="[5/6] Site image & Gantt...")
        if st.session_state.site_image_path:
            img_path = str(st.session_state.site_image_path)
            if os.path.exists(img_path): gen.insert_image(doc, img_path)

        # ── Cleanup + first save MUST happen before the Gantt XML surgery ──
        # Previously the Gantt block referenced `op` (the saved .docx path) while
        # `op` was only assigned further down — every run raised NameError and was
        # swallowed by 'Gantt insert skipped', so the Gantt never reached Word.
        prog.progress(80, text="[6/7] Page numbers & highlights...")
        gen.add_page_numbers(doc)
        gen.strip_all_highlights(doc)
        nm  = gen.highlight_missing(doc)
        ds  = datetime.now().strftime("%d_%b_%Y")
        sc  = re.sub(r"[^0-9A-Za-z]","",cap_s) or "XXX"
        scu = re.sub(r"[^0-9A-Za-z_ ]","",cust).replace(" ","_")
        ptype = "CAPEX" if is_capex else "OPEX"
        on  = f"Proposal_{ptype}_{scu}_{sc}kWp_{ds}.docx"
        op  = os.path.join(st.session_state.tmp_dir, on)
        doc.save(op)

        # ── Insert Gantt chart image into Word doc ──
        # NOTE: For OPEX, the "Tentative Project Schedule" table in the Word
        # template is left exactly as-is. Only the "Delivery & Installation"
        # placeholder (PH_100) is filled — via the normal gen.fill_document()
        # placeholder pass above — nothing else in that section is touched or
        # replaced with a generated image. The Gantt PNG generation/insertion
        # below runs for CAPEX only.
        prog.progress(88, text="[7/7] Gantt chart...")
        if is_capex:
            _timeline_gen = (st.session_state.get("f_PH_120") or st.session_state.get("f_PH_008") or
                             st.session_state.prefilled.get("PH_120","") or st.session_state.prefilled.get("PH_008","")).strip()
        else:
            _timeline_gen = ""

        # CAPEX Gantt data (from CAPEX Excel)
        # Same ground-truth bar data as the website preview (_GANTT_DATA above) —
        # read from the 4 reference sheets in Solar_Project_Gantt.xlsx.xlsx.
        _CAPEX_GANTT_GEN = {
            "3 + 1 Month": {"n_months":4,"bars":{1:[0],2:[0],3:[0],4:[0],5:[0],6:[0,1],7:[1],8:[1,2],9:[2,3],10:[3],11:[3]}},
            "4 + 1 Month": {"n_months":5,"bars":{1:[0],2:[0],3:[0],4:[0],5:[0,1],6:[1],7:[1],8:[2,3],9:[3,4],10:[4],11:[4]}},
            "5 + 1 Month": {"n_months":6,"bars":{1:[0],2:[0],3:[0],4:[0],5:[0,1],6:[1],7:[1],8:[2,3,4],9:[4,5],10:[5],11:[5]}},
            "6 + 1 Month": {"n_months":7,"bars":{1:[0],2:[0],3:[0],4:[0],5:[0,1],6:[1],7:[1,2],8:[2,3,4],9:[4,5,6],10:[4,5,6],11:[5,6]}},
        }
        _CAPEX_TASKS_GEN = [
            "Receipt of PO and Advance Order","Customer Kick-Off Meeting",
            "Design Approval / Site Handover (whichever later)","Submission & Approval of Design",
            "Placement of Orders on Vendors (Material Wise)","Manufacturing & Dispatch of Materials",
            "Mobilization of Team at Site","Installation & Test Charge","Approvals",
            "Commercial Operation Date (COD)","Site and Documentation Handover",
        ]

        # OPEX Gantt data (from Solar_Project_Gantt_-_OPEX_.xlsx)
        _OPEX_GANTT_GEN = {
            "4 + 1 Month": {
                "n_months":4,
                "tasks":["PPA Signing","Customer Kick-Off Meeting","Site Handover / Design Approval (whichever later)",
                         "Mobilization of team at site","Procurement","Civil Work, Installation & Cabling",
                         "Pre Commissioning & Test Charge","Liasoning","Site Acceptance Test","Handover","Project Closure"],
                "bars":{1:[0],2:[0],3:[0],4:[0],5:[0,1],6:[1,2],7:[2],8:[2,3],9:[2,3],10:[2,3],11:[]},
            },
            "5 + 1 Month": {
                "n_months":5,
                "tasks":["PPA Signing","Customer Kick-Off Meeting","Site Handover / Design Approval (whichever later)",
                         "Procurement","Mobilization of team at site","Civil Work, Installation & Cabling",
                         "Pre Commissioning & Test Charge","Liasoning","Site Acceptance Test","Handover","Project Closure"],
                "bars":{1:[0],2:[0],3:[0],4:[0],5:[0],6:[1,2,3],7:[2,3],8:[2,3],9:[3,4],10:[4],11:[]},
            },
            "6 + 1 Month": {
                "n_months":6,
                "tasks":["PPA Signing","Customer Kick-Off Meeting","Site Handover / Design Approval (whichever later)",
                         "Mobilization of team at site","Procurement","Civil Work, Installation & Cabling",
                         "Pre Commissioning & Test Charge","Liasoning","Site Acceptance Test","Handover","Project Closure"],
                "bars":{1:[0],2:[0],3:[0],4:[0],5:[0,1],6:[1],7:[1],8:[2,3,4],9:[4,5],10:[5],11:[]},
            },
        }

        if is_capex:
            _GANTT_DATA_GEN = _CAPEX_GANTT_GEN
            _TASKS_GEN_MAP  = {k: _CAPEX_TASKS_GEN for k in _CAPEX_GANTT_GEN}
        else:
            _GANTT_DATA_GEN = {k: {"n_months":v["n_months"],"bars":v["bars"]} for k,v in _OPEX_GANTT_GEN.items()}
            _TASKS_GEN_MAP  = {k: v["tasks"] for k,v in _OPEX_GANTT_GEN.items()}
        if _timeline_gen in _GANTT_DATA_GEN:
            try:
                from PIL import Image, ImageDraw, ImageFont
                from dateutil.relativedelta import relativedelta
                from datetime import date as _date
                import re as _re
                _gd        = _GANTT_DATA_GEN[_timeline_gen]
                _nm        = _gd["n_months"]
                _bars      = _gd["bars"]
                _TASKS_GEN = _TASKS_GEN_MAP[_timeline_gen]
                _today = _date.today()
                _months = [(_today + relativedelta(months=i+1)) for i in range(_nm)]
                _mo_labels = [m.strftime("%b-%y") for m in _months]
                _start_lbl = _months[0].strftime("%b-%y")
                _end_lbl   = _months[-1].strftime("%b-%y")

                # ── Dimensions matching Image 3 (Excel source) ──────────────
                # Total width = 640px → scales to 17cm in Word at 96dpi
                # Sr=4%, Task=38%, month cols split equally over remaining 58%
                # (Task column widened from 34% -> 38% to give the larger,
                # legible font enough room for the longest task descriptions.)
                _TOTAL_W  = 640
                _COL_SR   = int(_TOTAL_W * 0.04)          # ~26px
                _COL_TASK = int(_TOTAL_W * 0.38)          # ~243px
                _COL_MO   = (_TOTAL_W - _COL_SR - _COL_TASK) // _nm
                _ROW_H    = 22    # task row height (room for a legible font)
                _HDR_H    = 26    # header row height
                _W        = _COL_SR + _COL_TASK + _COL_MO * _nm
                _H        = _HDR_H + _ROW_H * 11 + _HDR_H  # header + 11 tasks + footer
                _SCALE    = 3    # 3x for crisp text
                img = Image.new("RGB", (_W*_SCALE, _H*_SCALE), "white")
                draw = ImageDraw.Draw(img)

                # Colors (exact from Image 3)
                _ORANGE = (233, 113, 50)
                _NAVY   = (31,  56,  100)
                _LGRAY  = (242, 242, 242)
                _WHITE  = (255, 255, 255)
                _BORDER = (217, 217, 217)

                def _rect(x, y, w, h, fill, outline=_BORDER, s=_SCALE):
                    draw.rectangle([x*s, y*s, (x+w)*s-1, (y+h)*s-1],
                                   fill=fill, outline=outline)

                # Try a handful of TrueType fonts across platforms — the previous
                # hardcoded Linux-only path silently fell back to PIL's tiny fixed
                # bitmap font (ImageFont.load_default(), which ignores fsize/scale
                # entirely) on Windows, which is why the text came out unreadably
                # small regardless of the fsize passed in.
                def _load_font(bold, size):
                    candidates = (
                        ["arialbd.ttf", "Arial Bold.ttf", "DejaVuSans-Bold.ttf",
                         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"] if bold else
                        ["arial.ttf", "Arial.ttf", "DejaVuSans.ttf",
                         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
                    )
                    for fp in candidates:
                        try:
                            return ImageFont.truetype(fp, size)
                        except Exception:
                            continue
                    return ImageFont.load_default()

                def _txt(x, y, w, h, text, fill=(0,0,0), align="left", s=_SCALE, bold=False, fsize=10):
                    font = _load_font(bold, fsize*s)
                    bbox = draw.textbbox((0,0), text, font=font)
                    tw, th = bbox[2]-bbox[0], bbox[3]-bbox[1]
                    if align == "center":
                        tx = x*s + (w*s - tw)//2
                    elif align == "right":
                        tx = x*s + w*s - tw - 3*s
                    else:
                        tx = x*s + 3*s
                    ty = y*s + (h*s - th)//2
                    draw.text((tx, ty), text, fill=fill, font=font)

                # ── Row 1: Column headers ──────────────────────────────────
                y = 0
                _rect(0,            y, _COL_SR,   _HDR_H, _NAVY, _WHITE)
                _txt( 0,            y, _COL_SR,   _HDR_H, "Sr.",             fill=_WHITE, align="center", bold=True)
                _rect(_COL_SR,      y, _COL_TASK, _HDR_H, _NAVY, _WHITE)
                _txt( _COL_SR,      y, _COL_TASK, _HDR_H, "Item Description",fill=_WHITE, align="left",   bold=True)
                for mi, lbl in enumerate(_mo_labels):
                    _x = _COL_SR + _COL_TASK + mi * _COL_MO
                    _rect(_x, y, _COL_MO, _HDR_H, _NAVY, _WHITE)
                    _txt( _x, y, _COL_MO, _HDR_H, lbl, fill=_WHITE, align="center", bold=True)

                # ── Task rows ─────────────────────────────────────────────
                for ti, task in enumerate(_TASKS_GEN):
                    y  = _HDR_H + ti * _ROW_H
                    bg = _LGRAY if ti % 2 == 0 else _WHITE
                    _rect(0,       y, _COL_SR,   _ROW_H, bg)
                    _txt( 0,       y, _COL_SR,   _ROW_H, str(ti+1), align="center", fsize=9)
                    _rect(_COL_SR, y, _COL_TASK, _ROW_H, bg)
                    _txt( _COL_SR, y, _COL_TASK, _ROW_H, task, align="left", fsize=9)
                    for mi in range(_nm):
                        _x = _COL_SR + _COL_TASK + mi * _COL_MO
                        filled = mi in _bars.get(ti+1, [])
                        _rect(_x, y, _COL_MO, _ROW_H, _ORANGE if filled else bg)

                # ── Footer — Image 3 style ─────────────────────────────────
                # Left block: Sr+Task cols → "Start Date ="  right-aligned
                # Month cols: first = start_lbl, middle empty navy, second-to-last = "Project End Date =", last = end_lbl
                y = _HDR_H + 11 * _ROW_H
                _rect(0, y, _COL_SR + _COL_TASK, _HDR_H, _NAVY, _WHITE)
                _txt( 0, y, _COL_SR + _COL_TASK, _HDR_H, "Start Date =",
                      fill=_WHITE, align="right", bold=True, fsize=9)
                for mi in range(_nm):
                    _x = _COL_SR + _COL_TASK + mi * _COL_MO
                    _rect(_x, y, _COL_MO, _HDR_H, _NAVY, _WHITE)
                    if mi == 0:
                        _txt(_x, y, _COL_MO, _HDR_H, _start_lbl,
                             fill=_WHITE, align="center", bold=True, fsize=9)
                    elif mi == _nm - 2 and _nm > 2:
                        _txt(_x, y, _COL_MO, _HDR_H, "Project End Date =",
                             fill=_WHITE, align="right", bold=True, fsize=8)
                    elif mi == _nm - 1:
                        _txt(_x, y, _COL_MO, _HDR_H, _end_lbl,
                             fill=_WHITE, align="center", bold=True, fsize=9)

                # Save PNG at full 3x resolution — do NOT downsize back to 1x.
                # (Previously the crisp 3x render was thrown away by resizing
                # back down to the tiny 1x canvas before saving, which is why
                # the text came out blurry/unreadably small in the Word doc.
                # Word displays this at a fixed width=Cm(17) regardless of
                # pixel size, so keeping the high-res version is what actually
                # makes the text sharp and legible at that print size.)
                _gantt_png = os.path.join(st.session_state.tmp_dir, "gantt_chart.png")
                img.save(_gantt_png, "PNG", dpi=(300,300))

                # ── Insert Gantt PNG into Word via python-docx (safe, cross-platform) ──
                # Previous approach: unzip → edit XML → rezip.
                # Problem on Windows: os.path.relpath uses backslash separators;
                # the .docx zip must use forward slashes throughout — Word rejected
                # the file with "experienced an error trying to open".
                # New approach: open the already-saved .docx with python-docx,
                # find the Gantt table by scanning paragraph text, replace it with
                # an inline image paragraph, then save again. No raw zip surgery.
                import copy
                from docx import Document as _Doc
                from docx.shared import Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from lxml import etree

                _gantt_doc = _Doc(op)

                # Locate the Gantt table: the first <w:tbl> that follows the
                # paragraph containing "Tentative Project Schedule" (skip TOC —
                # we want the second occurrence).
                _body = _gantt_doc.element.body
                _children = list(_body)
                _schedule_hits = []
                for _i, _child in enumerate(_children):
                    if _child.tag == qn('w:p'):
                        _txt = "".join(
                            r.text or "" for r in _child.iter(qn('w:t'))
                        )
                        if 'Tentative Project Schedule' in _txt:
                            _schedule_hits.append(_i)

                _gantt_tbl_idx = None
                # Search after the second hit (or first if only one)
                _search_from = _schedule_hits[-1] if _schedule_hits else 0
                for _i in range(_search_from + 1, len(_children)):
                    if _children[_i].tag == qn('w:tbl'):
                        _gantt_tbl_idx = _i
                        break

                if _gantt_tbl_idx is not None:
                    # Build a centred paragraph with the Gantt image
                    _new_para = _gantt_doc.add_paragraph()
                    _new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run = _new_para.add_run()
                    _run.add_picture(_gantt_png, width=Cm(17))   # 17 cm ≈ A4 usable width
                    # Detach the paragraph from the end of the document body
                    _new_para_elem = _new_para._element
                    _body.remove(_new_para_elem)
                    # Replace the Gantt table with the image paragraph
                    _old_tbl = _children[_gantt_tbl_idx]
                    _body.replace(_old_tbl, _new_para_elem)
                else:
                    # No table found — append the image after the heading
                    _new_para = _gantt_doc.add_paragraph()
                    _new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run = _new_para.add_run()
                    _run.add_picture(_gantt_png, width=Cm(17))

                _gantt_doc.save(op)   # overwrite in place — no new file, no rename

            except Exception as _ge:
                import traceback
                st.warning(f"Gantt insert skipped — document generated without it: {_ge}")
                with st.expander("Gantt error detail"): st.code(traceback.format_exc())
        prog.progress(100, text="Done!")
        st.session_state.output_path=op; st.session_state.output_name=on; st.session_state.n_missing=nm
        _pf_log = st.session_state.get("prefilled", {})
        # "Submitted by" is a DD_SUBMITTER widget — its selected value lives in the
        # widget's own session key (f_PH_016), not in prefilled, until write-out time.
        _submitter_raw = str(st.session_state.get("f_PH_016") or _pf_log.get("PH_016","") or "")
        _submitter_name = _submitter_raw.split("|")[0].strip() if _submitter_raw else ""
        _cap_log = _pf_log.get("PH_001") or st.session_state.get("capacity_kwp","")
        _started_at = st.session_state.get("_proposal_started_at")
        _elapsed_str = format_elapsed(datetime.now().timestamp() - _started_at) if _started_at else "—"
        st.session_state["_last_time_taken"] = _elapsed_str
        log_proposal(_submitter_name, _cap_log, st.session_state.get("project_type",""), on, _elapsed_str)
        st.session_state.stage="done"; st.rerun()
    except Exception as e:
        prog.progress(100, text="Error!"); st.error(f"Failed: {e}")
        import traceback
        with st.expander("Full error"): st.code(traceback.format_exc())
        if st.button("← Back"): st.session_state.stage="form"; st.rerun()


# ═══════════════════════════════ DONE ══════════════════════════
elif st.session_state.stage == "done":
    show_header("done")
    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,var(--cyan-lt),#F3FBFE);
      border:1.5px solid var(--cyan-md);border-radius:16px;
      padding:1.5rem 1.75rem;margin:.5rem 0 1.25rem;
      box-shadow:0 6px 20px var(--cyan-glow);">
      <div style="font-family:'Nunito',sans-serif;font-size:1.6rem;font-weight:900;
        color:var(--grey-dk);letter-spacing:-.01em;">✅ Proposal ready!</div>
      <div style="font-size:1.05rem;font-weight:800;color:var(--cyan-dk);margin-top:.3rem;
        font-family:'Nunito',sans-serif;word-break:break-all;">
        {st.session_state.output_name}</div>
    </div>""", unsafe_allow_html=True)
    if st.session_state.n_missing: st.warning(f"{st.session_state.n_missing} blank input(s) marked 'XX' in red.")
    c1,c2,c3,c4 = st.columns(4)
    with c1:
        with open(st.session_state.output_path,"rb") as f:
            st.download_button("📥 Download Word", data=f.read(),
                file_name=st.session_state.output_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary", use_container_width=True)
    with c2:
        with open(st.session_state.final_input_path,"rb") as f:
            st.download_button("📥 Download Excel", data=f.read(),
                file_name="Proposal_Input_Updated.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.document",
                use_container_width=True)
    with c3:
        if st.button("✏️ Edit & Re-download", use_container_width=True):
            # Go back to form — all session state preserved, user edits and regenerates
            st.session_state.stage = "form"
            st.rerun()
    with c4:
        if st.button("🔄 New Proposal", use_container_width=True):
            for k in list(st.session_state.keys()): del st.session_state[k]
            st.rerun()

    # ── Log box: every proposal generated so far ──
    st.markdown('<hr class="oline">', unsafe_allow_html=True)
    st.markdown('<div class="sec">Proposal History</div>', unsafe_allow_html=True)
    _all_log_rows = read_proposal_log(limit=None)
    _total_count = len(_all_log_rows)
    _total_cap = 0.0
    for _r in _all_log_rows:
        try:
            _total_cap += float(str(_r.get("Capacity (kWp)","") or "").replace(",","").replace("kWp","").strip())
        except Exception:
            pass

    # ── White stat box: total proposals made + total capacity across all of them ──
    st.markdown(f"""
    <div style="display:flex;gap:1rem;margin:.75rem 0 1.25rem;">
      <div style="flex:1;background:var(--white);border:1.5px solid var(--border2);border-radius:14px;padding:1.1rem 1.5rem;
        box-shadow:var(--shadow-sm);">
        <div style="font-family:'Nunito',sans-serif;font-size:1.9rem;font-weight:900;color:var(--grey-dk) !important;letter-spacing:-.02em;">
          {_total_count}</div>
        <div style="font-size:.75rem;font-weight:800;color:var(--orange) !important;margin-top:2px;">
          Total Proposals Made</div>
      </div>
      <div style="flex:1;background:var(--white);border:1.5px solid var(--border2);border-radius:14px;padding:1.1rem 1.5rem;
        box-shadow:var(--shadow-sm);">
        <div style="font-family:'Nunito',sans-serif;font-size:1.9rem;font-weight:900;color:var(--grey-dk) !important;letter-spacing:-.02em;">
          {fmt_indian(int(round(_total_cap)))} kWp</div>
        <div style="font-size:.75rem;font-weight:800;color:var(--orange) !important;margin-top:2px;">
          Total Capacity Proposed</div>
      </div>
    </div>""", unsafe_allow_html=True)

    _log_rows = _all_log_rows[:25]
    if not _log_rows:
        st.caption("No past proposals logged yet.")
    else:
        _rows_html = ""
        for r in _log_rows:
            _rows_html += f"""<tr>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.8rem;color:var(--text3);">{r.get('Date','')}</td>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.8rem;font-weight:700;color:var(--grey-dk);">{r.get('Submitted By','')}</td>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.8rem;color:var(--grey-dk);">{r.get('Capacity (kWp)','')} kWp</td>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.8rem;color:var(--text3);">{r.get('Project Type','')}</td>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.75rem;color:var(--text3);word-break:break-all;">{r.get('File Name','')}</td>
              <td style="padding:.5rem .75rem;border-bottom:1px solid var(--border2);font-size:.8rem;font-weight:700;color:var(--orange-dk);white-space:nowrap;">{r.get('Time Taken','—')}</td>
            </tr>"""
        st.markdown(f"""
        <div style="background:var(--white);border:1px solid var(--border);border-radius:12px;
          overflow-x:auto;box-shadow:var(--shadow-sm);">
          <table style="width:100%;border-collapse:collapse;font-family:'Nunito Sans',sans-serif;">
            <thead>
              <tr style="background:var(--grey-lt);">
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">Date</th>
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">Submitted By</th>
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">Capacity</th>
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">Type</th>
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">File</th>
                <th style="padding:.5rem .75rem;text-align:left;font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--grey);">Time Taken</th>
              </tr>
            </thead>
            <tbody>{_rows_html}</tbody>
          </table>
        </div>""", unsafe_allow_html=True)

# ═══════════════════════════════ FOOTER ══════════════════════════
# Rendered last so it always appears at the bottom of the page,
# in normal document flow (no position:fixed — see the earlier "!"
# help-button issue where guessed fixed-positioning CSS on a mismatched
# selector pushed the element off-screen; this avoids that entirely).
st.markdown("""
<style>
.app-footer {
  position: fixed;
  left: 0;
  right: 0;
  bottom: 0;
  z-index: 999;
  margin-top: 2.5rem;
  padding: 1rem 1.5rem;
  border-top: 1px solid var(--border);
  background: var(--white);
  display: flex;
  align-items: center;
  justify-content: center;
  flex-wrap: wrap;
  gap: .55rem;
  font-family: 'Nunito Sans', 'Inter', sans-serif;
}
.main .block-container {
  padding-bottom: 5rem;
}
.app-footer .f-item {
  font-weight: 700;
  font-size: .78rem;
  color: #3A3A3C;               /* dark grey */
  letter-spacing: .01em;
}
.app-footer .f-item.f-primary {
  font-weight: 800;
  color: #2C2C2E;                /* darker grey — top of hierarchy */
}
.app-footer .f-sep {
  color: #B0B0B3;
  font-weight: 400;
  padding: 0 .15rem;
}
.app-footer .f-contact {
  color: #58595B;
  font-weight: 600;
}
</style>
<div class="app-footer">
  <span class="f-item f-primary">Proposal Automator V1.0</span>
  <span class="f-sep">|</span>
  <span class="f-item">Developed by Carl Aaron</span>
  <span class="f-sep">|</span>
  <span class="f-item">Business Development</span>
  <span class="f-sep">|</span>
  <span class="f-item f-contact">Support &amp; Queries: +91 95510 62519</span>
</div>
""", unsafe_allow_html=True)
